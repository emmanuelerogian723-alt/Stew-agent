"""
S.T.E.W Book Generator — Professional long-form book creation engine.
Generates books up to 200 pages with front/back cover design.
"""
import io
import os
import re
import time
import logging
import requests
from urllib.parse import quote

logger = logging.getLogger(__name__)

# ── COVER IMAGE GENERATION (Pollinations) ──────────────────────────────────

def _fetch_cover_image(prompt: str, width: int = 1024, height: int = 1536,
                       max_retries: int = 3) -> bytes | None:
    """Fetch a book cover image from Pollinations AI."""
    url = f"https://image.pollinations.ai/prompt/{quote(prompt)}"
    params = {"width": width, "height": height, "model": "flux", "nologo": "true", "seed": 42}
    for attempt in range(max_retries):
        try:
            r = requests.get(url, params=params, timeout=90)
            if r.status_code == 200 and len(r.content) > 5000:
                return r.content
            logger.warning(f"Cover image attempt {attempt+1}: status={r.status_code}, size={len(r.content)}")
        except Exception as e:
            logger.warning(f"Cover image attempt {attempt+1} error: {e}")
        time.sleep(3 * (attempt + 1))
    return None


def _generate_cover_prompt(title: str, genre: str, is_back: bool = False) -> str:
    """Build an AI image prompt for a professional book cover."""
    if is_back:
        return (
            f"Professional book back cover design, {genre} genre, "
            f"dark elegant background with subtle texture, "
            f"space for author bio and barcode, "
            f"matching the front cover style of '{title}', "
            f"minimalist, high quality book design, 4k"
        )
    styles = {
        "fiction": "dramatic cinematic lighting, compelling character silhouette",
        "nonfiction": "clean professional typography layout, abstract geometric design",
        "business": "corporate professional design, gold and navy blue accents",
        "self_help": "inspirational sunrise motif, warm uplifting colors",
        "religion": "sacred majestic light from above, peaceful serene atmosphere",
        "children": "colorful playful illustration, whimsical cartoon style",
        "romance": "soft romantic pastel colors, elegant floral motifs",
        "thriller": "dark mysterious atmosphere, dramatic shadows, suspenseful",
        "biography": "dignified portrait-style composition, warm sepia tones",
        "academic": "scholarly clean design, university press style",
        "fantasy": "epic fantasy landscape, magical glowing elements",
        "scifi": "futuristic sci-fi design, neon and dark space theme",
    }
    style_desc = styles.get(genre, styles["nonfiction"])
    return (
        f"Professional book front cover for '{title}', {genre} genre, "
        f"{style_desc}, book title prominently displayed, author name at bottom, "
        f"publishing quality, 4k high resolution, portrait orientation"
    )


# ── BOOK OUTLINE GENERATION ─────────────────────────────────────────────────

def _safe_content(result) -> str:
    """Defensively extract text content from an LLM callback result.
    Handles: dict {"content": str}, dict {"content": dict} (double-wrap bug),
    raw str, or anything else — always returns a str, never crashes on slicing."""
    if isinstance(result, dict):
        inner = result.get("content", "")
        if isinstance(inner, dict):
            inner = inner.get("content", "")
        return inner if isinstance(inner, str) else str(inner) if inner else ""
    if isinstance(result, str):
        return result
    return str(result) if result else ""


def build_book_outline_prompt(topic: str, num_chapters: int, pages: int) -> str:
    """System prompt for generating a book outline."""
    return (
        f"You are a professional book author and editor. Create a detailed book outline "
        f"for a {pages}-page book about '{topic}'. The book should have {num_chapters} chapters. "
        f"Return ONLY a JSON array of objects, each with 'title' (chapter title) and "
        f"'summary' (2-3 sentence description of what the chapter covers). "
        f"The first chapter should be an introduction/overview. "
        f"The last chapter should be a conclusion. "
        f"Make chapter titles compelling and professional. JSON array only."
    )


def build_chapter_content_prompt(topic: str, chapter_title: str, chapter_summary: str,
                                  chapter_num: int, total_chapters: int, pages_per_chapter: int) -> str:
    """System prompt for generating a single chapter's content."""
    return (
        f"You are a professional book author. Write Chapter {chapter_num} of {total_chapters} "
        f"for a book about '{topic}'.\n\n"
        f"Chapter title: {chapter_title}\n"
        f"Chapter summary: {chapter_summary}\n\n"
        f"Write approximately {pages_per_chapter} pages of rich, engaging, professional content. "
        f"Use clear section headings (prefixed with '## '), sub-headings (prefixed with '### '), "
        f"and well-structured paragraphs. Include examples, anecdotes, and practical insights. "
        f"Write in a flowing narrative style appropriate for a published book.\n\n"
        f"DO NOT include the chapter number or title at the top — just start with the first '## ' section. "
        f"DO NOT use markdown bold (**text**) or italic (*text*) markers. "
        f"End with a brief transition to the next chapter."
    )


# ── BOOK COMPILATION (DOCX with covers) ────────────────────────────────────

def generate_book_docx(chapters: list[dict], title: str, author: str,
                       genre: str, front_cover_bytes: bytes | None = None,
                       back_cover_bytes: bytes | None = None) -> dict:
    """Generate a professional book DOCX with front cover, TOC, chapters, back cover."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import tempfile, base64

    doc = Document()

    # Set page size to 6x9 (trade paperback)
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # ── FRONT COVER ─────────────────────────────────────────────────────
    if front_cover_bytes:
        try:
            cover_path = tempfile.mktemp(suffix=".jpg")
            with open(cover_path, "wb") as f:
                f.write(front_cover_bytes)
            doc.add_picture(cover_path, width=Inches(4.5))
            # Center the cover image
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            os.unlink(cover_path)
        except Exception as e:
            logger.warning(f"Could not embed front cover: {e}")
            # Text-only cover fallback
            cover_title = doc.add_paragraph()
            cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cover_title.add_run(title)
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            doc.add_paragraph()
    else:
        cover_title = doc.add_paragraph()
        cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cover_title.add_run(title)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        if author:
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            arun = author_para.add_run(f"by {author}")
            arun.font.size = Pt(14)
            arun.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()

    doc.add_page_break()

    # ── COPYRIGHT PAGE ──────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    copyright_p = doc.add_paragraph()
    copyright_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = copyright_p.add_run(f"Copyright (c) {time.strftime('%Y')} {author or 'Author'}")
    crun.font.size = Pt(10)
    crun.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    cp2 = doc.add_paragraph()
    cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp2.add_run("All rights reserved.\n").font.size = Pt(10)
    cp2.add_run("Generated by S.T.E.W Agent").font.size = Pt(9)

    doc.add_page_break()

    # ── TABLE OF CONTENTS ───────────────────────────────────────────────
    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for i, ch in enumerate(chapters, 1):
        toc_p = doc.add_paragraph()
        toc_p.paragraph_format.space_after = Pt(4)
        run = toc_p.add_run(f"Chapter {i}: {ch.get('title', f'Chapter {i}')}")
        run.font.size = Pt(12)
        # Page number placeholder
        toc_p.add_run(f"  ....  {i * (200 // max(len(chapters), 1))}").font.size = Pt(10)

    doc.add_page_break()

    # ── CHAPTERS ───────────────────────────────────────────────────────
    for i, ch in enumerate(chapters, 1):
        # Chapter title page
        ch_title_p = doc.add_paragraph()
        ch_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ch_title_p.paragraph_format.space_before = Inches(2)
        ch_num_run = ch_title_p.add_run(f"Chapter {i}")
        ch_num_run.font.size = Pt(14)
        ch_num_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        ch_name_p = doc.add_paragraph()
        ch_name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ch_name_run = ch_name_p.add_run(ch.get("title", f"Chapter {i}"))
        ch_name_run.font.size = Pt(22)
        ch_name_run.font.bold = True
        ch_name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        doc.add_paragraph()

        # Chapter content
        content = ch.get("content", "")
        lines = content.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
            elif stripped.startswith("## "):
                heading_text = stripped[3:].replace("**", "").replace("*", "")
                h = doc.add_heading(heading_text, level=2)
            elif stripped.startswith("### "):
                heading_text = stripped[4:].replace("**", "").replace("*", "")
                h = doc.add_heading(heading_text, level=3)
            elif stripped.startswith("# "):
                heading_text = stripped[2:].replace("**", "").replace("*", "")
                h = doc.add_heading(heading_text, level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                clean = stripped.lstrip("-* ").replace("**", "").replace("*", "")
                doc.add_paragraph(clean, style="List Bullet")
            else:
                p = doc.add_paragraph(stripped.replace("**", "").replace("*", ""))
                p.paragraph_format.first_line_indent = Inches(0.3)
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.line_spacing = 1.15

        # Page break after each chapter (except last)
        if i < len(chapters):
            doc.add_page_break()

    # ── BACK COVER ──────────────────────────────────────────────────────
    doc.add_page_break()
    if back_cover_bytes:
        try:
            back_path = tempfile.mktemp(suffix=".jpg")
            with open(back_path, "wb") as f:
                f.write(back_cover_bytes)
            doc.add_picture(back_path, width=Inches(4.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            os.unlink(back_path)
        except Exception as e:
            logger.warning(f"Could not embed back cover: {e}")
            blurb = doc.add_paragraph()
            blurb.alignment = WD_ALIGN_PARAGRAPH.CENTER
            blurb.add_run(f"About {author}\n").font.bold = True
            blurb.add_run(f"{author} is an author and creator. This book was generated with S.T.E.W Agent.")

    # ── SAVE ────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    safe_title = re.sub(r'[^\w\- ]', '', title)[:60].strip().replace(" ", "_") or "book"
    filename = f"{safe_title}.docx"

    return {
        "file_bytes": buf.getvalue(),
        "filename": filename,
        "format": "docx",
        "pages_estimated": sum(len(ch.get("content", "")) // 3000 for ch in chapters),
        "chapters": len(chapters),
    }


# ── BOOK GENERATION ORCHESTRATOR ────────────────────────────────────────────

def generate_book(topic: str, author: str = "", pages: int = 100,
                   llm_chat_fn=None, llm_complete_fn=None) -> dict:
    """Full book generation pipeline.
    
    Args:
        topic: Book subject/title
        author: Author name
        pages: Target page count (up to 200)
        llm_chat_fn: async fn(messages, max_tokens) -> {"content": str}
        llm_complete_fn: sync fn(prompt, system) -> str
    
    Returns dict with file_bytes, filename, etc.
    """
    import json as _json

    pages = max(10, min(pages, 200))

    # Determine genre from topic
    topic_lower = topic.lower()
    genre = "nonfiction"
    genre_map = {
        "business": "business", "entrepreneur": "business", "startup": "business",
        "money": "business", "finance": "business", "invest": "business",
        "church": "religion", "god": "religion", "faith": "religion",
        "christian": "religion", "prayer": "religion", "spiritual": "religion",
        "love": "romance", "romance": "romance", "relationship": "romance",
        "child": "children", "kids": "children",
        "biography": "biography", "memoir": "biography", "life story": "biography",
        "novel": "fiction", "story": "fiction", "fiction": "fiction",
        "fantasy": "fantasy", "magic": "fantasy",
        "science fiction": "scifi", "sci-fi": "scifi", "space": "scifi",
        "thriller": "thriller", "mystery": "thriller", "crime": "thriller",
        "study": "academic", "research": "academic", "textbook": "academic",
        "self help": "self_help", "motivation": "self_help", "success": "self_help",
    }
    for keyword, g in genre_map.items():
        if keyword in topic_lower:
            genre = g
            break

    # Calculate chapters
    num_chapters = max(5, min(pages // 10, 20))
    pages_per_chapter = pages // num_chapters

    # Step 1: Generate outline
    logger.info(f"Generating book outline: {num_chapters} chapters, {pages} pages, genre={genre}")
    outline_prompt = build_book_outline_prompt(topic, num_chapters, pages)
    outline_msg = [{"role": "system", "content": outline_prompt},
                   {"role": "user", "content": f"Book topic: {topic}"}]

    outline_raw = None
    if llm_chat_fn:
        result = llm_chat_fn(outline_msg, max_tokens=4000)
        outline_raw = _safe_content(result)
    elif llm_complete_fn:
        outline_raw = llm_complete_fn(f"Book topic: {topic}", system=outline_prompt)

    chapters = []
    if outline_raw:
        json_match = re.search(r'\[.*\]', outline_raw, re.DOTALL)
        if json_match:
            try:
                chapters = _json.loads(json_match.group())
            except Exception:
                pass

    if not chapters:
        # Fallback: generate generic chapters
        chapters = [{"title": f"Chapter {i}: {topic}", "summary": f"Content about {topic}"} 
                     for i in range(1, num_chapters + 1)]

    # Step 2: Generate cover images (non-blocking, done first)
    logger.info("Generating book cover images...")
    front_prompt = _generate_cover_prompt(topic, genre, is_back=False)
    back_prompt = _generate_cover_prompt(topic, genre, is_back=True)
    
    front_cover = _fetch_cover_image(front_prompt, width=768, height=1152)
    back_cover = _fetch_cover_image(back_prompt, width=768, height=1152)

    # Step 3: Generate each chapter's content
    for i, ch in enumerate(chapters):
        ch_title = ch.get("title", f"Chapter {i+1}")
        ch_summary = ch.get("summary", "")
        
        logger.info(f"Generating chapter {i+1}/{len(chapters)}: {ch_title}")
        chapter_sys = build_chapter_content_prompt(
            topic, ch_title, ch_summary, i + 1, len(chapters), pages_per_chapter
        )
        
        chapter_content = None
        if llm_chat_fn:
            result = llm_chat_fn(
                [{"role": "system", "content": chapter_sys},
                 {"role": "user", "content": f"Write chapter {i+1} now."}],
                max_tokens=8000
            )
            chapter_content = _safe_content(result)
        elif llm_complete_fn:
            chapter_content = llm_complete_fn(f"Write chapter {i+1} now.", system=chapter_sys)

        if not chapter_content:
            chapter_content = f"## {ch_title}\n\nContent for this chapter about {topic}."

        ch["content"] = chapter_content

    # Step 4: Compile the book
    logger.info("Compiling book DOCX...")
    return generate_book_docx(chapters, topic, author, genre, front_cover, back_cover)


# ── SONG / MUSIC GENERATION ──────────────────────────────────────────────────
#
# Free/open-source engine stack (no paid GPU hosting required):
#   1. ACE-Step 1.5 (primary) — real lyric-based singing generation, running on
#      HuggingFace's free ZeroGPU pool via the public Space's Gradio API.
#      This is the closest free/open-source equivalent to Suno/Gemini music —
#      it actually sings the lyrics, not just instrumental. Best quality, but
#      shared community GPU queue means it can take 30s-3min depending on load.
#   2. Meta MusicGen-small (fallback) — HuggingFace serverless Inference API,
#      instrumental-only, no vocals, but reliably fast.
#   3. Pollinations TTS (last resort) — spoken-word reading of the lyrics over
#      a synthesized voice, so the user always gets *something* even if both
#      GPU-backed engines are down.
# Lyrics/title/genre/mood: generated via whatever LLM Stew already has
# configured (Groq Llama 3, Mistral, etc. — see llm_client.py fallback chain).
# Cover art: FLUX.1 via Pollinations (see _fetch_cover_image, already FLUX).

GENRE_TAGS = {
    "afrobeat": "afrobeat, log drum, percussion, guitar, horns, groovy, upbeat, 106 bpm",
    "gospel": "gospel, choir, piano, organ, uplifting, soulful, powerful vocals, 90 bpm",
    "hip-hop": "hip hop, 808 bass, trap drums, boom bap, confident, rhythmic, 95 bpm",
    "hip hop": "hip hop, 808 bass, trap drums, boom bap, confident, rhythmic, 95 bpm",
    "rap": "rap, hard hitting drums, deep 808 bass, aggressive, rhythmic flow, 90 bpm",
    "drill": "drill, sliding 808s, hi-hats, dark atmosphere, gritty, 140 bpm",
    "amapiano": "amapiano, log drum, piano, shaker, deep house groove, smooth, 112 bpm",
    "highlife": "highlife, guitar, horns, percussion, danceable, warm, joyful, 118 bpm",
    "fuji": "fuji, talking drum, percussion, vocal chant, rhythmic, traditional, 120 bpm",
    "r&b": "r&b, smooth bass, soulful vocals, mellow, romantic, 80 bpm",
    "rnb": "r&b, smooth bass, soulful vocals, mellow, romantic, 80 bpm",
    "pop": "pop, catchy melody, synth, upbeat, polished production, 118 bpm",
    "rock": "rock, electric guitar, drums, bass, powerful, energetic, 130 bpm",
    "country": "country, acoustic guitar, banjo, warm, storytelling, twangy, 100 bpm",
    "jazz": "jazz, saxophone, double bass, piano, swing, smooth, improvisational, 100 bpm",
    "blues": "blues, electric guitar, slow groove, soulful, raw emotion, 70 bpm",
    "classical": "classical, orchestral strings, piano, elegant, cinematic, 90 bpm",
    "orchestra": "orchestral, full symphony, strings, brass, epic, cinematic, 90 bpm",
    "edm": "edm, synth lead, four on the floor kick, energetic, festival, 128 bpm",
    "house": "house music, four on the floor, synth bass, groovy, hypnotic, 124 bpm",
    "techno": "techno, driving kick, minimal synth, hypnotic, dark, 130 bpm",
}


def _detect_genre(text: str) -> str:
    """Scan free-text for a known genre keyword. Returns '' if none found."""
    lower = text.lower()
    for genre in GENRE_TAGS:
        if genre in lower:
            return genre
    return ""


def _generate_ace_step_song(tags: str, lyrics: str, duration: float = 60.0,
                             hf_token: str = "") -> tuple:
    """Generate a real sung song (vocals + instrumentation) via the free,
    open-source ACE-Step 1.5 model, called through its public HuggingFace
    Space Gradio API (runs on HF's shared ZeroGPU pool — free, but queued).
    Returns (audio_bytes, format_ext) or (None, None) on failure/timeout."""
    try:
        from gradio_client import Client
        import concurrent.futures

        def _call():
            # NOTE: gradio_client's Client takes `token`, not `hf_token`.
            # ACE-Step's public Space works fine fully anonymous (no token needed).
            client = Client("ACE-Step/ACE-Step", token=hf_token or None)
            result = client.predict(
                audio_duration=duration,
                prompt=tags,
                lyrics=lyrics,
                infer_step=60,
                guidance_scale=15.0,
                scheduler_type="euler",
                cfg_type="apg",
                omega_scale=10.0,
                manual_seeds=str(int(time.time()) % 2_000_000_000),
                guidance_interval=0.5,
                guidance_interval_decay=0.0,
                min_guidance_scale=3.0,
                use_erg_tag=True,
                use_erg_lyric=True,
                use_erg_diffusion=True,
                oss_steps="",
                guidance_scale_text=0.0,
                guidance_scale_lyric=0.0,
                audio2audio_enable=False,
                ref_audio_strength=0.5,
                ref_audio_input=None,
                lora_name_or_path="none",
                api_name="/__call__",
            )
            audio_path = result[0] if isinstance(result, (list, tuple)) else result
            if audio_path and os.path.exists(audio_path):
                ext = os.path.splitext(audio_path)[1].lstrip(".").lower() or "mp3"
                with open(audio_path, "rb") as f:
                    return (f.read(), ext)
            return (None, None)

        # ZeroGPU is a shared community queue — bound the wait so we can fall
        # back to MusicGen/TTS instead of leaving the user hanging forever.
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
            future = ex.submit(_call)
            return future.result(timeout=200)
    except concurrent.futures.TimeoutError:
        logger.warning("ACE-Step generation timed out (ZeroGPU queue busy) — falling back")
        return (None, None)
    except Exception as e:
        logger.warning(f"ACE-Step generation failed: {e}")
        return (None, None)


def generate_song(prompt: str, llm_complete_fn=None, llm_chat_fn=None,
                   duration_seconds: int = 60, genre_hint: str = "") -> dict:
    """Generate a complete song: title, genre, mood, structured lyrics
    (verse/chorus/bridge/outro), album cover art, and real audio.

    Engine order: ACE-Step 1.5 (real singing) -> MusicGen (instrumental)
    -> Pollinations TTS (spoken lyrics fallback).

    Returns dict with audio_bytes, lyrics, cover_bytes, filename, title,
    genre, mood, engine_used.
    """
    import json as _json

    genre = genre_hint or _detect_genre(prompt)

    # Step 1: Generate structured song data (title, genre, mood, tags, lyrics) via LLM
    songwriter_prompt = (
        "You are a professional songwriter and music producer. Write an original song "
        "based on the user's request. Return ONLY a JSON object with these exact keys:\n"
        '"title" (a short catchy song title), '
        '"genre" (one word/phrase best describing the musical genre), '
        '"mood" (one word describing the emotional mood, e.g. joyful, melancholic, energetic), '
        '"tags" (comma-separated music style descriptors for an AI music generator — instruments, '
        'tempo in BPM, genre, mood — e.g. "afrobeat, log drum, guitar, percussion, upbeat, 106 bpm"), '
        '"lyrics" (the full song lyrics using [Verse], [Chorus], [Bridge], [Outro] section tags, '
        "each on its own line before that section's lyrics; include at least 2 verses, a chorus "
        "repeated between verses, a bridge, and an outro).\n"
        + (f"The requested genre is {genre}. " if genre else "")
        + "Make it catchy, emotional, and memorable. JSON only, no markdown, no extra text."
    )

    raw = ""
    if llm_chat_fn:
        result = llm_chat_fn(
            [{"role": "system", "content": songwriter_prompt},
             {"role": "user", "content": f"Write a song about: {prompt}"}],
            max_tokens=2500
        )
        raw = _safe_content(result)
    elif llm_complete_fn:
        raw = llm_complete_fn(f"Write a song about: {prompt}", system=songwriter_prompt)

    title, mood, tags, lyrics = "", "", "", ""
    if raw:
        json_match = re.search(r'\{.*\}', raw, re.DOTALL)
        if json_match:
            try:
                data = _json.loads(json_match.group())
                title = data.get("title", "")
                genre = genre or data.get("genre", "")
                mood = data.get("mood", "")
                tags = data.get("tags", "")
                lyrics = data.get("lyrics", "")
            except Exception as e:
                logger.warning(f"Song JSON parse failed: {e}")

    if not lyrics:
        # Fallback: plain lyrics without structured JSON (older model behavior)
        lyrics = (f"[Verse]\n{prompt}\n\n[Chorus]\n{prompt}\n\n"
                  f"[Verse]\n{prompt}\n\n[Bridge]\n{prompt}\n\n[Outro]\n{prompt}")
    if not title:
        title = prompt[:60].strip().title() or "Untitled Song"
    if not tags:
        tags = GENRE_TAGS.get(genre, f"{genre or 'pop'}, melodic, upbeat, 110 bpm")
    if not genre:
        genre = "pop"

    # Step 2: Generate album cover (FLUX.1 via Pollinations)
    cover_prompt = (
        f"Professional album cover art for a {genre} song titled '{title}' about "
        f"'{prompt[:80]}', {mood or 'vibrant'} mood, musical theme, high quality, "
        f"square format, 4k"
    )
    cover_bytes = _fetch_cover_image(cover_prompt, width=1024, height=1024)

    # Step 3: Generate audio — try engines in order of quality
    audio_bytes = None
    audio_format = "wav"
    engine_used = "none"

    hf_token = os.environ.get("HF_TOKEN") or os.environ.get("HUGGINGFACE_API_KEY") or ""

    # Engine 1: ACE-Step 1.5 — real singing with the actual lyrics (free ZeroGPU)
    logger.info(f"Attempting song generation via ACE-Step 1.5 (genre={genre})...")
    ace_bytes, ace_ext = _generate_ace_step_song(tags, lyrics, duration=min(max(duration_seconds, 30), 90),
                                                  hf_token=hf_token)
    if ace_bytes and len(ace_bytes) > 1000:
        audio_bytes = ace_bytes
        audio_format = ace_ext or "mp3"
        engine_used = "ace-step-1.5"
        logger.info(f"ACE-Step succeeded: {len(audio_bytes)} bytes ({audio_format})")

    # Engine 2: MusicGen-small — instrumental only, but fast and reliable
    if not audio_bytes and hf_token:
        try:
            logger.info("Falling back to MusicGen-small (instrumental)...")
            hf_url = "https://api-inference.huggingface.co/models/facebook/musicgen-small"
            headers = {"Authorization": f"Bearer {hf_token}"}
            payload = {
                "inputs": f"{tags}, high quality music, professional production",
                "parameters": {
                    "do_sample": True,
                    "max_new_tokens": min(duration_seconds * 50, 1500),
                    "temperature": 1.0,
                    "guidance_scale": 3.0,
                }
            }
            for attempt in range(3):
                resp = requests.post(hf_url, headers=headers, json=payload, timeout=120)
                if resp.status_code == 200:
                    audio_bytes = resp.content
                    audio_format = "wav"
                    engine_used = "musicgen-small"
                    logger.info(f"MusicGen returned {len(audio_bytes)} bytes of audio")
                    break
                elif resp.status_code == 503:
                    logger.info(f"MusicGen loading (503), retry {attempt+1}...")
                    time.sleep(15)
                else:
                    logger.warning(f"MusicGen API error: {resp.status_code} - {resp.text[:200]}")
                    break
        except Exception as e:
            logger.warning(f"MusicGen generation failed: {e}")

    # Engine 3: Pollinations TTS — spoken-word lyrics, last resort.
    # Pollinations deprecated the free anonymous audio endpoint (text.pollinations.ai
    # with model=openai-audio now 404s); audio now lives at gen.pollinations.ai and
    # requires an API key (POLLINATIONS_API_KEY). Skip cleanly if no key is configured
    # instead of burning retries on a guaranteed 401.
    if not audio_bytes:
        polli_key = os.environ.get("POLLINATIONS_API_KEY", "")
        if polli_key:
            logger.info("Falling back to Pollinations TTS for spoken lyrics...")
            try:
                tts_text = lyrics[:2000]
                resp = requests.get(
                    f"https://gen.pollinations.ai/audio/{quote(tts_text)}",
                    params={"voice": "nova"},
                    headers={"Authorization": f"Bearer {polli_key}"},
                    timeout=90,
                )
                if resp.status_code == 200 and len(resp.content) > 1000:
                    audio_bytes = resp.content
                    audio_format = "mp3"
                    engine_used = "tts-fallback"
                    logger.info(f"TTS returned {len(audio_bytes)} bytes")
                else:
                    logger.warning(f"TTS fallback: status={resp.status_code}")
            except Exception as e:
                logger.warning(f"TTS fallback failed: {e}")
        else:
            logger.info("Pollinations TTS skipped — no POLLINATIONS_API_KEY configured "
                        "(their free anonymous audio tier was discontinued)")

    safe_name = re.sub(r'[^\w\- ]', '', title)[:40].strip().replace(" ", "_") or "song"

    return {
        "title": title,
        "genre": genre,
        "mood": mood,
        "lyrics": lyrics,
        "cover_bytes": cover_bytes,
        "audio_bytes": audio_bytes,
        "audio_format": audio_format,
        "filename": f"{safe_name}.{audio_format}",
        "engine_used": engine_used,
    }
