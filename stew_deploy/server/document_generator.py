"""
S.T.E.W Document Generator — REAL binary file generation.
Supports: PDF, DOCX, XLSX, PPTX, HTML
All files returned as base64-encoded strings.
"""
import base64
import io
import json
import logging
from datetime import datetime
from typing import Any, Optional

from fastapi import HTTPException

logger = logging.getLogger(__name__)


def _to_base64(buf: io.BytesIO) -> str:
    buf.seek(0)
    return base64.b64encode(buf.read()).decode("utf-8")


# Common unicode chars that break in reportlab's default Helvetica font
_UNICODE_REPLACEMENTS = {
    "\u2248": "~",     # ≈
    "\u00d7": "x",     # ×
    "\u2212": "-",     # − (unicode minus)
    "\u2013": "-",     # – (en dash)
    "\u2014": "--",    # — (em dash)
    "\u2018": "'", "\u2019": "'",  # smart quotes
    "\u201c": '"', "\u201d": '"',
    "\u2192": "->",    # →
    "\u03bb": "lambda", "\u039b": "Lambda",
    "\u03c1": "rho", "\u03b4": "delta",
    "\u03bc": "u", "\u00b5": "u",  # micro/mu
    "\u2103": " deg C",
    "\u00b1": "+/-",   # ±
    "\u2265": ">=", "\u2264": "<=",
    "\u2260": "!=",
    "\u221e": "infinity",
    "\u00b0": " deg",
    "\u25a0": "",      # black square (the ■ artifact itself)
    "\ufffd": "",      # replacement character
}

# Superscript/subscript digit maps -> plain ASCII (e.g. 10⁻³² -> 10^-32)
_SUPERSCRIPT_MAP = str.maketrans({
    "\u2070": "0", "\u00b9": "1", "\u00b2": "2", "\u00b3": "3",
    "\u2074": "4", "\u2075": "5", "\u2076": "6", "\u2077": "7",
    "\u2078": "8", "\u2079": "9", "\u207b": "-",
})
_SUBSCRIPT_MAP = str.maketrans({
    "\u2080": "0", "\u2081": "1", "\u2082": "2", "\u2083": "3",
    "\u2084": "4", "\u2085": "5", "\u2086": "6", "\u2087": "7",
    "\u2088": "8", "\u2089": "9", "\u208b": "-",
})


# Currency symbols to preserve (replace with ASCII-safe equivalents before stripping)
_CURRENCY_PRESERVE = {
    "\u20a6": "NGN ",   # ₦ Naira
    "\u20ac": "EUR ",  # € Euro
    "\u00a3": "GBP ",  # £ Pound
    "\u00a2": " cents", # ¢ cent
    "\u00a5": "JPY ",  # ¥ Yen
    "\u20b9": "INR ",  # ₹ Rupee
    "\u20a9": "KRW ",  # ₩ Won
    "\u00ab": "<<",    # «
    "\u00bb": ">>",    # »
    "\u2026": "...",   # … ellipsis
    "\u2022": "*",     # • bullet (for PDF safety)
    "\u00a0": " ",     # non-breaking space
    "\u00a9": "(c)",   # ©
    "\u00ae": "(R)",   # ®
    "\u2122": "(TM)",  # ™
    "\u00b7": "-",     # · middle dot
    "\u00d7": "x",     # × multiply (also in _UNICODE_REPLACEMENTS but ensure)
    "\u00f7": "/",     # ÷ divide
    "\u2192": "->",    # →
    "\u2190": "<-",    # ←
    "\u2191": "^",     # ↑
    "\u2193": "v",     # ↓
    "\u2713": "OK",    # ✓ check
    "\u2717": "X",     # ✗
    "\u2605": "*",     # ★ star
    "\u2606": "*",     # ☆
}


def _sanitize_text(text: str) -> str:
    """Strip/replace unicode chars that render as black boxes in default PDF/DOCX fonts.
    Preserves currency symbols (incl. Naira) by converting to ASCII equivalents first."""
    if not text:
        return text
    text = text.translate(_SUPERSCRIPT_MAP)
    text = text.translate(_SUBSCRIPT_MAP)
    # Preserve currency symbols and common chars before stripping
    for uni, replacement in _CURRENCY_PRESERVE.items():
        text = text.replace(uni, replacement)
    for uni, replacement in _UNICODE_REPLACEMENTS.items():
        text = text.replace(uni, replacement)
    # Strip any remaining non-ASCII characters (final safety net)
    text = text.encode("ascii", "ignore").decode("ascii")
    return text


def _strip_markdown(text: str) -> str:
    """Remove ALL markdown formatting markers from text for clean document output.
    Strips: ##, ###, **, *, __, _, `, ---, >, and bullet markers.
    Used for slide content, table cells, and any text that should be plain."""
    if not text:
        return text
    import re as _re
    # Remove heading markers: ## Heading -> Heading
    text = _re.sub(r'^#{1,6}\s+', '', text, flags=_re.MULTILINE)
    # Remove bold: **text** or __text__ -> text
    text = _re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = _re.sub(r'__(.+?)__', r'\1', text)
    # Remove italic: *text* or _text_ -> text (but not list bullets or math)
    text = _re.sub(r'(?<!\w)\*([^*\n]+?)\*(?!\w)', r'\1', text)
    text = _re.sub(r'(?<!\w)_([^_\n]+?)_(?!\w)', r'\1', text)
    # Remove inline code backticks
    text = _re.sub(r'`([^`\n]+?)`', r'\1', text)
    # Remove horizontal rules
    text = _re.sub(r'^[\s]*[-_]{3,}[\s]*$', '', text, flags=_re.MULTILINE)
    # Remove blockquotes
    text = _re.sub(r'^>\s+', '', text, flags=_re.MULTILINE)
    return text.strip()


# ── PDF ───────────────────────────────────────────────────────────────────────

def _parse_md_table_row(line: str) -> list:
    """Parse a markdown table row '| a | b | c |' into ['a','b','c']."""
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return [_sanitize_text(_strip_markdown(c)) for c in cells]


def _is_table_separator(line: str) -> bool:
    """Detect '|---|---|---|' style separator rows."""
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    inner = stripped.strip("|")
    return all(c in "-:| " for c in inner) and "-" in inner


def generate_pdf(content: str, title: str = "Document") -> dict:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
            title=_sanitize_text(title),
            author="S.T.E.W Agent",
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Title"],
            fontSize=18,
            spaceAfter=12,
            alignment=TA_CENTER,
        )
        body_style = ParagraphStyle(
            "CustomBody",
            parent=styles["Normal"],
            fontSize=11,
            leading=16,
            spaceAfter=8,
        )
        table_style = TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563EB")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EFF6FF")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ])

        story = []

        # Cover Page
        story.append(Spacer(1, 5 * cm))
        story.append(Paragraph(_sanitize_text(title), ParagraphStyle(
            "Cover", parent=styles["Title"], fontSize=28, leading=34, spaceAfter=20,
            alignment=TA_CENTER, textColor=colors.HexColor("#0F172A"), fontName="Helvetica-Bold",
        )))
        story.append(Spacer(1, 0.8 * cm))
        story.append(Paragraph("Generated by S.T.E.W Agent", ParagraphStyle(
            "CoverSub", parent=styles["Normal"], fontSize=12, alignment=TA_CENTER,
            textColor=colors.HexColor("#64748B"),
        )))
        story.append(Paragraph(datetime.utcnow().strftime("%B %d, %Y"), ParagraphStyle(
            "CoverDate", parent=styles["Normal"], fontSize=12, alignment=TA_CENTER,
            textColor=colors.HexColor("#64748B"),
        )))
        from reportlab.platypus import PageBreak
        story.append(PageBreak())

        lines = content.split("\n")
        i = 0
        n = len(lines)
        while i < n:
            line = lines[i]
            stripped = line.strip()

            if stripped == "<!--PAGEBREAK-->":
                story.append(PageBreak())
                i += 1
                continue

            # Detect start of a markdown table: a row with pipes, followed by a separator row
            if stripped.startswith("|") and i + 1 < n and _is_table_separator(lines[i + 1]):
                header_row = _parse_md_table_row(stripped)
                table_data = [header_row]
                j = i + 2
                while j < n and lines[j].strip().startswith("|"):
                    table_data.append(_parse_md_table_row(lines[j]))
                    j += 1
                # Wrap cell text in Paragraphs so long text wraps instead of overflowing
                cell_style = ParagraphStyle("Cell", parent=styles["Normal"], fontSize=9, leading=11)
                wrapped = [[Paragraph(str(cell), cell_style) for cell in row] for row in table_data]
                col_count = len(header_row) or 1
                avail_width = 17 * cm
                col_width = avail_width / col_count
                t = Table(wrapped, colWidths=[col_width] * col_count, repeatRows=1)
                t.setStyle(table_style)
                story.append(t)
                story.append(Spacer(1, 0.4 * cm))
                i = j
                continue

            if not stripped:
                story.append(Spacer(1, 0.3 * cm))
            elif stripped.startswith("### "):
                clean_h = _sanitize_text(_strip_markdown(stripped[4:]))
                story.append(Paragraph(clean_h, ParagraphStyle("H3", parent=styles["Heading3"],
                    fontSize=12, textColor=colors.HexColor("#475569"), spaceBefore=10, spaceAfter=4)))
            elif stripped.startswith("## "):
                clean_h = _sanitize_text(_strip_markdown(stripped[3:]))
                story.append(Paragraph(clean_h, ParagraphStyle("H2", parent=styles["Heading2"],
                    fontSize=13, textColor=colors.HexColor("#334155"), spaceBefore=12, spaceAfter=6)))
            elif stripped.startswith("# "):
                clean_h = _sanitize_text(_strip_markdown(stripped[2:]))
                story.append(Paragraph(clean_h, ParagraphStyle("H1", parent=styles["Heading1"],
                    fontSize=16, textColor=colors.HexColor("#1E3A5F"), spaceBefore=16, spaceAfter=8)))
            elif stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("• "):
                bullet_text = _sanitize_text(_strip_markdown(stripped.lstrip("-*• ")))
                story.append(Paragraph(f"• {bullet_text}", body_style))
            else:
                # Strip ALL markdown markers, sanitize unicode, escape HTML chars
                safe = _sanitize_text(_strip_markdown(stripped))
                safe = safe.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, body_style))
            i += 1

        # Footer with timestamp
        story.append(Spacer(1, 1 * cm))
        footer_style = ParagraphStyle("Footer", parent=styles["Normal"], fontSize=8,
                                      textColor=colors.grey, alignment=TA_CENTER)
        story.append(Paragraph(
            f"Generated by S.T.E.W Agent - {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
            footer_style,
        ))

        def _page_footer(canv, doc_):
            canv.saveState()
            canv.setFont("Helvetica", 8)
            canv.setFillColor(colors.HexColor("#94A3B8"))
            page_num = canv.getPageNumber()
            if page_num > 1:
                canv.drawCentredString(A4[0] / 2, 1.2 * cm, f"Page {page_num}")
            canv.drawRightString(A4[0] - 2 * cm, 1.2 * cm, "S.T.E.W Agent")
            canv.restoreState()
        doc.build(story, onFirstPage=_page_footer, onLaterPages=_page_footer)
        filename = f"{title.replace(' ', '_')}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.pdf"
        return {
            "file": _to_base64(buf),
            "filename": filename,
            "mime_type": "application/pdf",
            "success": True,
        }
    except ImportError:
        raise HTTPException(500, "reportlab not installed — add it to requirements.txt")
    except Exception as e:
        logger.error(f"PDF generation error: {e}")
        raise HTTPException(500, f"PDF generation failed: {e}")


# ── DOCX ──────────────────────────────────────────────────────────────────────

def generate_docx(content: str, title: str = "Document") -> dict:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        # Set heading colors
        for heading_name, color in [("Heading 1", RGBColor(0x1E, 0x3A, 0x5F)), ("Heading 2", RGBColor(0x33, 0x41, 0x55)), ("Heading 3", RGBColor(0x47, 0x55, 0x69))]:
            try:
                doc.styles[heading_name].font.color.rgb = color
                doc.styles[heading_name].font.name = "Calibri"
            except Exception:
                pass

        # Title page
        title_para = doc.add_heading(_sanitize_text(title), level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            run.font.size = Pt(28)

        # Subtitle
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run("Generated by S.T.E.W Agent")
        sub_run.font.size = Pt(12)
        sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        sub_run.font.italic = True

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.utcnow().strftime("%B %d, %Y"))
        date_run.font.size = Pt(11)
        date_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_page_break()
        doc.add_paragraph()  # spacer

        def _shade_cell(cell, color_hex):
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), color_hex)
            cell._tc.get_or_add_tcPr().append(shading)

        lines = content.split("\n")
        i = 0
        n = len(lines)
        while i < n:
            line = lines[i]
            stripped = line.strip()

            # Detect markdown table
            if stripped.startswith("|") and i + 1 < n and _is_table_separator(lines[i + 1]):
                header_row = _parse_md_table_row(stripped)
                table_rows = []
                j = i + 2
                while j < n and lines[j].strip().startswith("|"):
                    table_rows.append(_parse_md_table_row(lines[j]))
                    j += 1

                col_count = len(header_row) or 1
                table = doc.add_table(rows=1, cols=col_count)
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for idx, htext in enumerate(header_row):
                    hdr_cells[idx].text = htext
                    _shade_cell(hdr_cells[idx], "2563EB")
                    for p in hdr_cells[idx].paragraphs:
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                for row_idx, row_data in enumerate(table_rows):
                    row_cells = table.add_row().cells
                    for idx in range(col_count):
                        val = row_data[idx] if idx < len(row_data) else ""
                        row_cells[idx].text = val
                        if row_idx % 2 == 1:
                            _shade_cell(row_cells[idx], "EFF6FF")
                doc.add_paragraph()  # spacer after table
                i = j
                continue

            if stripped == "<!--PAGEBREAK-->":
                doc.add_page_break()
                i += 1
                continue

            if not stripped:
                doc.add_paragraph()
            elif stripped.startswith("### "):
                doc.add_heading(_sanitize_text(_strip_markdown(stripped[4:])), level=3)
            elif stripped.startswith("## "):
                doc.add_heading(_sanitize_text(_strip_markdown(stripped[3:])), level=2)
            elif stripped.startswith("# "):
                doc.add_heading(_sanitize_text(_strip_markdown(stripped[2:])), level=1)
            elif stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("• "):
                clean_bullet = _sanitize_text(_strip_markdown(stripped.lstrip("-*• ")))
                doc.add_paragraph(clean_bullet, style="List Bullet")
            elif stripped.startswith("1. ") or (len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == "."):
                doc.add_paragraph(_sanitize_text(_strip_markdown(stripped)), style="List Number")
            else:
                doc.add_paragraph(_sanitize_text(_strip_markdown(stripped)))
            i += 1

        # Footer paragraph
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(
            f"Generated by S.T.E.W Agent - {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        buf = io.BytesIO()
        doc.save(buf)
        filename = f"{title.replace(' ', '_')}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.docx"
        return {
            "file": _to_base64(buf),
            "filename": filename,
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "success": True,
        }
    except ImportError:
        raise HTTPException(500, "python-docx not installed")
    except Exception as e:
        logger.error(f"DOCX generation error: {e}")
        raise HTTPException(500, f"DOCX generation failed: {e}")


# ── XLSX ──────────────────────────────────────────────────────────────────────

def generate_xlsx(
    data: list[dict],
    sheet_name: str = "Sheet1",
    title: str = "Spreadsheet",
) -> dict:
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name[:31]

        if not data:
            data = [{"No data": "provided"}]

        headers = list(data[0].keys())

        # Title row (row 1)
        title_font = Font(color="FFFFFF", bold=True, size=14)
        title_fill = PatternFill(start_color="1E3A5F", end_color="1E3A5F", fill_type="solid")
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
        title_cell = ws.cell(row=1, column=1, value=title)
        title_cell.font = title_font
        title_cell.fill = title_fill
        title_cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 32

        # Subtitle row (row 2)
        sub_font = Font(color="64748B", italic=True, size=10)
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))
        sub_cell = ws.cell(row=2, column=1, value=f"Generated by S.T.E.W Agent | {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')} | {len(data)} rows")
        sub_cell.font = sub_font
        sub_cell.alignment = Alignment(horizontal="center")
        ws.row_dimensions[2].height = 20

        # Header row (row 3)
        HEADER_ROW = 3
        header_fill = PatternFill(start_color="2563EB", end_color="2563EB", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=11)
        border = Border(
            left=Side(style="thin", color="CBD5E1"),
            right=Side(style="thin", color="CBD5E1"),
            top=Side(style="thin", color="CBD5E1"),
            bottom=Side(style="thin", color="CBD5E1"),
        )

        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=HEADER_ROW, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border

        # Data rows
        alt_fill = PatternFill(start_color="F1F5F9", end_color="F1F5F9", fill_type="solid")
        for row_idx, row in enumerate(data, HEADER_ROW + 1):
            for col_idx, header in enumerate(headers, 1):
                val = row.get(header, "")
                if isinstance(val, str) and val.replace(",", "").replace(".", "").replace("-", "").isdigit():
                    try:
                        val = float(val) if "." in val else int(val.replace(",", ""))
                    except ValueError:
                        pass
                cell = ws.cell(row=row_idx, column=col_idx, value=val)
                cell.border = border
                cell.alignment = Alignment(vertical="center")
                if (row_idx - HEADER_ROW) % 2 == 0:
                    cell.fill = alt_fill
                if isinstance(val, (int, float)) and not isinstance(val, bool):
                    cell.number_format = "#,##0" if isinstance(val, int) else "#,##0.00"

        # Auto-fit columns
        for col_idx, header in enumerate(headers, 1):
            max_len = max(
                len(str(header)),
                *[len(str(row.get(header, ""))) for row in data],
            )
            ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 50)

        ws.freeze_panes = f"A{HEADER_ROW + 1}"
        ws.auto_filter.ref = f"A{HEADER_ROW}:{get_column_letter(len(headers))}{HEADER_ROW + len(data)}"

        buf = io.BytesIO()
        wb.save(buf)
        filename = f"{title.replace(' ', '_')}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.xlsx"
        return {
            "file": _to_base64(buf),
            "filename": filename,
            "mime_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "success": True,
        }
    except ImportError:
        raise HTTPException(500, "openpyxl not installed")
    except Exception as e:
        logger.error(f"XLSX generation error: {e}")
        raise HTTPException(500, f"XLSX generation failed: {e}")




# ── PPTX ──────────────────────────────────────────────────────────────────────

def generate_pptx(slides: list[dict], title: str = "Presentation", theme: str = None,
                   use_images: bool = True) -> dict:
    """Generate a premium PPTX with 50+ professional themes.
    Theme is auto-detected from the title/topic, or can be specified by name.
    When use_images is True (default), fetches a real AI-generated hero photo
    for the title and closing slides.
    Image fetch failures are silent."""
    try:
        from pptx import Presentation
        from pptx.util import Inches, Pt, Emu
        from pptx.dml.color import RGBColor
        from pptx.enum.text import PP_ALIGN, MSO_ANCHOR
        from pptx.enum.shapes import MSO_SHAPE
        from pptx.oxml.ns import qn
        from server.slide_themes import render_pptx, auto_select_theme, THEMES, _fetch_hero_image

        prs = Presentation()
        SLIDE_W = 13.333
        SLIDE_H = 7.5
        prs.slide_width = Inches(SLIDE_W)
        prs.slide_height = Inches(SLIDE_H)

        if not theme:
            theme = auto_select_theme(title)

        category = THEMES.get(theme, {}).get("category", "corporate")

        hero_image = None
        closing_image = None
        if use_images:
            import time as _time
            hero_image = _fetch_hero_image(title, category, seed=1)
            if len(slides) > 3:
                _time.sleep(1.5)
                closing_image = _fetch_hero_image(title, category, seed=2)

        render_pptx(prs, slides, title, theme_name=theme, hero_image=hero_image, closing_image=closing_image)

        buf = io.BytesIO()
        prs.save(buf)
        clean_title = title.replace(" ", "_").replace("/", "_")[:40]
        filename = f"{clean_title}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.pptx"
        return {
            "file": _to_base64(buf),
            "filename": filename,
            "mime_type": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            "success": True,
            "theme": theme,
            "theme_category": category,
            "available_themes": len(THEMES),
            "hero_image_used": hero_image is not None,
            "closing_image_used": closing_image is not None,
        }
    except ImportError:
        raise HTTPException(500, "python-pptx not installed")
    except Exception as e:
        logger.error(f"PPTX generation error: {e}")
        raise HTTPException(500, f"PPTX generation failed: {e}")

# ── HTML ──────────────────────────────────────────────────────────────────────

def generate_html(content: str, title: str = "Report") -> dict:
    import html as html_module

    # Convert markdown-ish content to HTML
    lines = content.split("\n")
    body_html = []
    in_ul = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if in_ul:
                body_html.append("</ul>")
                in_ul = False
            body_html.append("<br>")
        elif stripped.startswith("## "):
            if in_ul:
                body_html.append("</ul>")
                in_ul = False
            body_html.append(f"<h2>{html_module.escape(stripped[3:])}</h2>")
        elif stripped.startswith("# "):
            if in_ul:
                body_html.append("</ul>")
                in_ul = False
            body_html.append(f"<h1>{html_module.escape(stripped[2:])}</h1>")
        elif stripped.startswith("### "):
            if in_ul:
                body_html.append("</ul>")
                in_ul = False
            body_html.append(f"<h3>{html_module.escape(stripped[4:])}</h3>")
        elif stripped.startswith("- ") or stripped.startswith("* "):
            if not in_ul:
                body_html.append("<ul>")
                in_ul = True
            body_html.append(f"<li>{html_module.escape(stripped[2:])}</li>")
        else:
            if in_ul:
                body_html.append("</ul>")
                in_ul = False
            body_html.append(f"<p>{html_module.escape(stripped)}</p>")

    if in_ul:
        body_html.append("</ul>")

    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{html_module.escape(title)}</title>
  <style>
    *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{
      font-family: 'Segoe UI', system-ui, -apple-system, sans-serif;
      background: #0f172a;
      color: #e2e8f0;
      line-height: 1.7;
      padding: 2rem;
    }}
    .container {{ max-width: 900px; margin: 0 auto; }}
    header {{
      background: linear-gradient(135deg, #1e293b, #0f172a);
      border-left: 4px solid #38bdf8;
      padding: 1.5rem 2rem;
      margin-bottom: 2rem;
      border-radius: 0 8px 8px 0;
    }}
    header h1 {{ color: #38bdf8; font-size: 2rem; }}
    header small {{ color: #94a3b8; font-size: 0.85rem; }}
    .content {{ background: #1e293b; padding: 2rem; border-radius: 8px; }}
    h1, h2, h3 {{ color: #38bdf8; margin: 1.5rem 0 0.75rem; }}
    h1 {{ font-size: 1.8rem; }}
    h2 {{ font-size: 1.4rem; border-bottom: 1px solid #334155; padding-bottom: 0.4rem; }}
    h3 {{ font-size: 1.1rem; color: #7dd3fc; }}
    p {{ margin-bottom: 0.8rem; color: #cbd5e1; }}
    ul {{ padding-left: 1.5rem; margin-bottom: 0.8rem; }}
    li {{ margin-bottom: 0.4rem; color: #cbd5e1; }}
    footer {{
      text-align: center;
      margin-top: 2rem;
      color: #64748b;
      font-size: 0.8rem;
    }}
    @media (max-width: 600px) {{ body {{ padding: 1rem; }} }}
  </style>
</head>
<body>
  <div class="container">
    <header>
      <h1>{html_module.escape(title)}</h1>
      <small>Generated by S.T.E.W Agent • {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}</small>
    </header>
    <div class="content">
      {''.join(body_html)}
    </div>
    <footer>S.T.E.W — Structured Task Execution Workflow</footer>
  </div>
</body>
</html>"""

    buf = io.BytesIO(html_content.encode("utf-8"))
    filename = f"{title.replace(' ', '_')}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.html"
    return {
        "file": _to_base64(buf),
        "filename": filename,
        "mime_type": "text/html",
        "success": True,
    }


# ── TERM PAPER / PRESENTATION PDF ──────────────────────────────────────────────

def generate_term_paper_pdf(
    content: str,
    title: str = "Term Paper",
    university: str = "University of Nigeria, Nsukka",
    department: str = "",
    author: str = "",
    reg_no: str = "",
    level: str = "",
    course_code: str = "",
    course_title: str = "",
    lecturer: str = "",
    paper_date: str = "",
    doc_type_label: str = "A TERM PAPER ON",
) -> dict:
    """Generate a strict academic term paper / presentation PDF following
    the UNN format pattern. Designed for students who need professional
    presentation documents with a proper cover page, table of contents,
    numbered sections, and references.

    Follows any user-provided details (university, department, course,
    lecturer, reg number, etc.) to customize the document.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm, mm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak, Flowable, KeepTogether
        )
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        import re as _re

        buf = io.BytesIO()
        PAGE_W, PAGE_H = A4
        LEFT_M = 3 * cm
        RIGHT_M = 2.5 * cm
        TOP_M = 2.5 * cm
        BOTTOM_M = 2.5 * cm
        CONTENT_W = PAGE_W - LEFT_M - RIGHT_M

        # Date formatting
        if not paper_date:
            paper_date = datetime.utcnow().strftime("%B %d, %Y")

        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            rightMargin=RIGHT_M,
            leftMargin=LEFT_M,
            topMargin=TOP_M,
            bottomMargin=BOTTOM_M,
            title=_sanitize_text(title),
            author=f"S.T.E.W Agent - {author}" if author else "S.T.E.W Agent",
        )

        styles = getSampleStyleSheet()

        # ── Cover page styles (centered, clean academic) ──
        cover_uni_style = ParagraphStyle(
            "CoverUni", parent=styles["Normal"], fontSize=14, leading=18,
            alignment=TA_CENTER, textColor=colors.HexColor("#1a1a1a"),
            fontName="Helvetica-Bold", spaceAfter=4,
        )
        cover_dept_style = ParagraphStyle(
            "CoverDept", parent=styles["Normal"], fontSize=13, leading=16,
            alignment=TA_CENTER, textColor=colors.HexColor("#333333"),
            fontName="Helvetica-Bold", spaceAfter=20,
        )
        cover_label_style = ParagraphStyle(
            "CoverLabel", parent=styles["Normal"], fontSize=12, leading=16,
            alignment=TA_CENTER, textColor=colors.HexColor("#444444"),
            fontName="Helvetica", spaceAfter=6,
        )
        cover_title_style = ParagraphStyle(
            "CoverTitle", parent=styles["Title"], fontSize=15, leading=20,
            alignment=TA_CENTER, textColor=colors.HexColor("#1a1a1a"),
            fontName="Helvetica-Bold", spaceAfter=20,
        )
        cover_field_style = ParagraphStyle(
            "CoverField", parent=styles["Normal"], fontSize=12, leading=18,
            alignment=TA_CENTER, textColor=colors.HexColor("#333333"),
            fontName="Helvetica", spaceAfter=10,
        )
        cover_field_bold = ParagraphStyle(
            "CoverFieldBold", parent=styles["Normal"], fontSize=12, leading=18,
            alignment=TA_CENTER, textColor=colors.HexColor("#1a1a1a"),
            fontName="Helvetica-Bold", spaceAfter=10,
        )

        # ── Body styles ──
        toc_title_style = ParagraphStyle(
            "TOCTitle", parent=styles["Heading1"], fontSize=13, leading=16,
            textColor=colors.HexColor("#1a1a1a"), fontName="Helvetica-Bold",
            spaceAfter=16, alignment=TA_LEFT,
        )
        toc_entry_style = ParagraphStyle(
            "TOCEntry", parent=styles["Normal"], fontSize=12, leading=20,
            textColor=colors.HexColor("#333333"), spaceAfter=2,
            fontName="Helvetica",
        )
        toc_sub_entry_style = ParagraphStyle(
            "TOCSubEntry", parent=styles["Normal"], fontSize=11, leading=18,
            textColor=colors.HexColor("#555555"), spaceAfter=2,
            leftIndent=24, fontName="Helvetica",
        )
        section_heading_style = ParagraphStyle(
            "SectionHeading", parent=styles["Heading1"], fontSize=13, leading=16,
            textColor=colors.HexColor("#1a1a1a"), fontName="Helvetica-Bold",
            spaceBefore=16, spaceAfter=8, keepWithNext=True,
        )
        sub_heading_style = ParagraphStyle(
            "SubHeading", parent=styles["Heading2"], fontSize=12, leading=15,
            textColor=colors.HexColor("#333333"), fontName="Helvetica-Bold",
            spaceBefore=10, spaceAfter=6, keepWithNext=True,
        )
        body_text_style = ParagraphStyle(
            "BodyText", parent=styles["Normal"], fontSize=12, leading=18,
            textColor=colors.HexColor("#1a1a1a"), spaceAfter=8,
            alignment=TA_JUSTIFY, fontName="Helvetica",
            firstLineIndent=0,
        )
        bullet_style = ParagraphStyle(
            "Bullet", parent=body_text_style, leftIndent=20, bulletIndent=10,
            spaceAfter=4, alignment=TA_LEFT,
        )
        ref_style = ParagraphStyle(
            "RefStyle", parent=styles["Normal"], fontSize=11, leading=15,
            textColor=colors.HexColor("#333333"), spaceAfter=6,
            alignment=TA_LEFT, fontName="Helvetica",
            leftIndent=18, firstLineIndent=-18,  # hanging indent
        )
        ref_title_style = ParagraphStyle(
            "RefTitle", parent=styles["Heading1"], fontSize=13, leading=16,
            textColor=colors.HexColor("#1a1a1a"), fontName="Helvetica-Bold",
            spaceBefore=20, spaceAfter=12, alignment=TA_CENTER,
        )

        story = []

        # ═══════════════════════════════════════════════════════
        # COVER PAGE (First Layer — strict pattern)
        # ═══════════════════════════════════════════════════════
        story.append(Spacer(1, 3 * cm))

        # University name
        story.append(Paragraph(_sanitize_text(university.upper()), cover_uni_style))

        # Department
        if department:
            story.append(Paragraph(_sanitize_text(department.upper()), cover_dept_style))

        story.append(Spacer(1, 1 * cm))

        # "A TERM PAPER ON" label
        story.append(Paragraph(doc_type_label.upper(), cover_label_style))
        story.append(Spacer(1, 0.3 * cm))

        # Title (underlined)
        story.append(Paragraph(_sanitize_text(title), cover_title_style))

        story.append(Spacer(1, 1 * cm))

        # Presented by
        story.append(Paragraph("PRESENTED BY", cover_label_style))
        if author:
            story.append(Paragraph(_sanitize_text(author), cover_field_style))
        else:
            story.append(Paragraph("_______________________________________", cover_field_style))

        # Reg. No.
        story.append(Spacer(1, 0.3 * cm))
        if reg_no:
            story.append(Paragraph(f"REG. NO: {_sanitize_text(reg_no)}", cover_field_style))
        else:
            story.append(Paragraph("REG. NO: _______________________________", cover_field_style))

        # Level / Program
        if level:
            story.append(Paragraph(_sanitize_text(level.upper()), cover_field_style))

        # Course
        if course_code or course_title:
            course_line = ""
            if course_code:
                course_line = f"COURSE: {_sanitize_text(course_code)}"
            if course_title:
                if course_line:
                    course_line += f" \u2013 {_sanitize_text(course_title)}"
                else:
                    course_line = f"COURSE: {_sanitize_text(course_title)}"
            story.append(Paragraph(course_line, cover_field_style))

        # Lecturer
        if lecturer:
            story.append(Paragraph(f"LECTURER: {_sanitize_text(lecturer)}", cover_field_style))

        story.append(Spacer(1, 0.8 * cm))

        # Date
        story.append(Paragraph(_sanitize_text(paper_date), cover_field_style))

        story.append(PageBreak())

        # ═══════════════════════════════════════════════════════
        # TABLE OF CONTENTS (auto-generated from content)
        # ═══════════════════════════════════════════════════════
        lines = content.split("\n")
        toc_entries = []
        i = 0
        n = len(lines)

        # First pass: extract headings for TOC
        for line in lines:
            stripped = line.strip()
            # Match patterns like "1.0 Introduction" or "## 1.0 Introduction"
            # or "## Introduction" (markdown style)
            match = _re.match(r'^(?:#{1,3}\s+)?(\d+\.?\d*)\s+(.+)', stripped)
            if match:
                num = match.group(1)
                heading_text = _sanitize_text(_strip_markdown(match.group(2)))
                is_subsection = "." in num and num.split(".")[1] != "0"
                toc_entries.append((num, heading_text, is_subsection))
            elif stripped.lower() == "references":
                toc_entries.append(("", "References", False))

        story.append(Paragraph("TABLE OF CONTENTS", toc_title_style))
        story.append(Spacer(1, 0.3 * cm))

        for num, heading, is_sub in toc_entries:
            if is_sub:
                story.append(Paragraph(
                    f"{num} {heading}", toc_sub_entry_style
                ))
            else:
                label = f"{num} {heading}" if num else heading
                story.append(Paragraph(label, toc_entry_style))

        story.append(PageBreak())

        # ═══════════════════════════════════════════════════════
        # BODY CONTENT (Strict academic pattern)
        # ═══════════════════════════════════════════════════════
        in_references = False
        i = 0
        while i < n:
            line = lines[i]
            stripped = line.strip()

            if stripped == "<!--PAGEBREAK-->":
                story.append(PageBreak())
                i += 1
                continue

            # Check for References section
            if stripped.lower() in ("references", "## references", "# references"):
                in_references = True
                story.append(Paragraph("REFERENCES", ref_title_style))
                story.append(Spacer(1, 0.3 * cm))
                i += 1
                continue

            if in_references:
                # References: each line is a reference entry with hanging indent
                if not stripped:
                    story.append(Spacer(1, 0.15 * cm))
                    i += 1
                    continue
                ref_text = _sanitize_text(_strip_markdown(stripped))
                ref_text = ref_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(ref_text, ref_style))
                i += 1
                continue

            # Detect markdown table
            if stripped.startswith("|") and i + 1 < n and _is_table_separator(lines[i + 1]):
                header_row = _parse_md_table_row(stripped)
                table_data = [header_row]
                j = i + 2
                while j < n and lines[j].strip().startswith("|"):
                    table_data.append(_parse_md_table_row(lines[j]))
                    j += 1
                cell_style = ParagraphStyle("Cell", parent=styles["Normal"],
                                           fontSize=10, leading=12)
                wrapped = [[Paragraph(str(cell), cell_style) for cell in row]
                          for row in table_data]
                col_count = len(header_row) or 1
                col_width = CONTENT_W / col_count
                tbl = Table(wrapped, colWidths=[col_width] * col_count, repeatRows=1)
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a1a")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#999999")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                     [colors.white, colors.HexColor("#f5f5f5")]),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 0.3 * cm))
                i = j
                continue

            if not stripped:
                story.append(Spacer(1, 0.25 * cm))
            # Heading patterns: "## 1.0 Title" or "# 1.0 Title" or "1.0 Title"
            elif _re.match(r'^#{1,2}\s+\d+\.\d*\s+', stripped):
                # Section heading
                clean = _re.sub(r'^#{1,2}\s+', '', stripped)
                clean = _sanitize_text(_strip_markdown(clean))
                story.append(KeepTogether([
                    Paragraph(clean, section_heading_style),
                    Spacer(1, 0.1 * cm),
                ]))
            elif _re.match(r'^#{3}\s+\d+\.\d*\s+', stripped):
                # Subsection heading
                clean = _re.sub(r'^#{3}\s+', '', stripped)
                clean = _sanitize_text(_strip_markdown(clean))
                story.append(Paragraph(clean, sub_heading_style))
            elif _re.match(r'^\d+\.\d*\s+', stripped):
                # Numbered heading without markdown
                clean = _sanitize_text(_strip_markdown(stripped))
                story.append(KeepTogether([
                    Paragraph(clean, section_heading_style),
                    Spacer(1, 0.1 * cm),
                ]))
            elif stripped.startswith("### "):
                clean = _sanitize_text(_strip_markdown(stripped[4:]))
                story.append(Paragraph(clean, sub_heading_style))
            elif stripped.startswith("## "):
                clean = _sanitize_text(_strip_markdown(stripped[3:]))
                story.append(Paragraph(clean, section_heading_style))
            elif stripped.startswith("# "):
                clean = _sanitize_text(_strip_markdown(stripped[2:]))
                story.append(Paragraph(clean, section_heading_style))
            elif stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("\u2022 "):
                bullet_text = _sanitize_text(_strip_markdown(stripped.lstrip("-*\u2022 ")))
                safe = bullet_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(f"\u2022 {safe}", bullet_style))
            else:
                safe = _sanitize_text(_strip_markdown(stripped))
                safe = safe.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, body_text_style))
            i += 1

        # ═══════════════════════════════════════════════════════
        # PAGE FOOTER
        # ═══════════════════════════════════════════════════════
        def _cover_border(canv, doc_):
            """Draw a double-border box around the cover page, matching UNN format."""
            canv.saveState()
            canv.setStrokeColor(colors.HexColor("#1a1a1a"))
            outer_margin = 1.0 * cm
            inner_margin = 1.2 * cm
            canv.setLineWidth(1.4)
            canv.rect(outer_margin, outer_margin,
                      PAGE_W - 2 * outer_margin, PAGE_H - 2 * outer_margin)
            canv.setLineWidth(0.6)
            canv.rect(outer_margin + inner_margin * 0.25, outer_margin + inner_margin * 0.25,
                      PAGE_W - 2 * (outer_margin + inner_margin * 0.25),
                      PAGE_H - 2 * (outer_margin + inner_margin * 0.25))
            canv.restoreState()

        def _page_footer(canv, doc_):
            canv.saveState()
            page_num = canv.getPageNumber()
            if page_num > 1:
                canv.setFont("Helvetica", 8)
                canv.setFillColor(colors.HexColor("#888888"))
                # Page number centered
                canv.drawCentredString(PAGE_W / 2, 1.2 * cm, str(page_num))
            canv.restoreState()

        doc.build(story, onFirstPage=_cover_border, onLaterPages=_page_footer)

        filename = f"{title.replace(' ', '_')[:50]}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.pdf"
        return {
            "file": _to_base64(buf),
            "filename": filename,
            "mime_type": "application/pdf",
            "success": True,
            "doc_type": "term_paper",
        }
    except ImportError:
        raise HTTPException(500, "reportlab not installed")
    except Exception as e:
        logger.error(f"Term paper PDF generation error: {e}")
        raise HTTPException(500, f"Term paper PDF generation failed: {e}")

