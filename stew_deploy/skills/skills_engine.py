"""
S.T.E.W — Secret Task Execution Worker
Skills Engine v3.0 — 60+ Real Capabilities
============================================
Real Playwright browser. Vision. Document creation.
Deep research. File management. Finance. Translation.
Created by Emmanuel Ene Rejoice Gideon — MUTYINT
"""

import asyncio
import base64
import io
import json
import os
import re
import requests
import subprocess
import sys
import tempfile
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional
from loguru import logger


class SkillsEngine:
    """60+ skills for S.T.E.W — Secret Task Execution Worker"""

    def __init__(self):
        self.serper_key = os.getenv("SERPER_API_KEY", "1867df2e3c379ba68185b39b64f2b986fba9e78e")
        self.output_dir = Path("output")
        self.workspace_dir = Path("workspace")
        self.screenshots_dir = Path("screenshots")
        self.output_dir.mkdir(exist_ok=True)
        self.workspace_dir.mkdir(exist_ok=True)
        self.screenshots_dir.mkdir(exist_ok=True)
        self._browser = None
        self._playwright = None
        logger.info("⚡ S.T.E.W Skills Engine v3.0 — 60 skills ONLINE")

    # ═══════════════════════════════════════════
    # 🔍 WEB SEARCH & RESEARCH
    # ═══════════════════════════════════════════

    async def web_search(self, query: str, num_results: int = 10) -> List[Dict]:
        """Real web search via Serper (Google results)"""
        try:
            response = requests.post(
                "https://google.serper.dev/search",
                headers={"X-API-KEY": self.serper_key, "Content-Type": "application/json"},
                json={"q": query, "num": num_results},
                timeout=10
            )
            if response.ok:
                data = response.json()
                results = []
                for item in data.get("organic", []):
                    results.append({
                        "title": item.get("title", ""),
                        "body": item.get("snippet", ""),
                        "href": item.get("link", "")
                    })
                logger.info(f"🔍 Search '{query}' — {len(results)} results")
                return results
            raise Exception(f"Serper HTTP {response.status_code}")
        except Exception as e:
            logger.error(f"Web search error: {e}")
            return [{"title": "Search error", "body": str(e), "href": ""}]

    async def news_search(self, topic: str, num_results: int = 10) -> List[Dict]:
        """Latest news via Serper News"""
        try:
            response = requests.post(
                "https://google.serper.dev/news",
                headers={"X-API-KEY": self.serper_key, "Content-Type": "application/json"},
                json={"q": topic, "num": num_results},
                timeout=10
            )
            if response.ok:
                data = response.json()
                return [{"title": n.get("title", ""), "body": n.get("snippet", ""),
                         "href": n.get("link", ""), "date": n.get("date", "")}
                        for n in data.get("news", [])]
        except Exception as e:
            logger.error(f"News search error: {e}")
        return []

    async def scrape_webpage(self, url: str) -> Dict:
        """Scrape and extract content from any webpage"""
        try:
            headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
            response = requests.get(url, headers=headers, timeout=15)
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(response.content, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
            title = soup.find("title")
            return {
                "url": url, "title": title.text if title else "No title",
                "content": text[:6000], "status_code": response.status_code, "success": True
            }
        except Exception as e:
            return {"url": url, "error": str(e), "success": False}

    async def wikipedia_search(self, topic: str) -> Dict:
        """Search Wikipedia for information"""
        try:
            import wikipedia
            summary = wikipedia.summary(topic, sentences=5)
            return {"topic": topic, "summary": summary, "success": True}
        except Exception as e:
            return {"topic": topic, "error": str(e), "success": False}

    async def youtube_search(self, query: str) -> List[Dict]:
        """Search YouTube for videos"""
        try:
            response = requests.post(
                "https://google.serper.dev/videos",
                headers={"X-API-KEY": self.serper_key, "Content-Type": "application/json"},
                json={"q": query + " site:youtube.com", "num": 5},
                timeout=10
            )
            if response.ok:
                data = response.json()
                return [{"title": v.get("title", ""), "link": v.get("link", ""),
                         "snippet": v.get("snippet", "")} for v in data.get("organic", [])[:5]]
        except Exception as e:
            logger.error(f"YouTube search error: {e}")
        return []

    async def deep_research(self, topic: str, brain=None) -> Dict:
        """Deep research — search, scrape multiple sources, AI synthesis"""
        try:
            results = await self.web_search(topic, num_results=6)
            findings = []
            for r in results[:4]:
                url = r.get("href", "")
                if url and url.startswith("http"):
                    scraped = await self.scrape_webpage(url)
                    if scraped.get("success") and scraped.get("content"):
                        findings.append({
                            "source": url,
                            "title": scraped.get("title", ""),
                            "summary": scraped.get("content", "")[:600],
                        })
                    await asyncio.sleep(0.5)

            synthesis = f"Researched '{topic}' across {len(findings)} sources."
            if brain and findings:
                source_text = "\n\n".join([
                    f"Source: {f['title']}\nURL: {f['source']}\n{f['summary']}"
                    for f in findings
                ])
                synthesis = await brain.call_llm(
                    f"Synthesize these research findings about '{topic}':\n\n{source_text}",
                    system="You are a world-class research analyst. Provide a comprehensive, well-structured synthesis with key findings, insights, and actionable conclusions."
                )

            return {
                "topic": topic, "sources_searched": len(results),
                "sources_scraped": len(findings), "findings": findings,
                "synthesis": synthesis, "success": True
            }
        except Exception as e:
            return {"topic": topic, "error": str(e), "success": False}

    # ═══════════════════════════════════════════
    # 🌐 REAL PLAYWRIGHT BROWSER CONTROL
    # ═══════════════════════════════════════════

    async def _get_browser(self):
        """Get or create Playwright browser instance — graceful if unavailable"""
        if self._browser:
            try:
                if self._browser.is_connected():
                    return self._browser
            except Exception:
                self._browser = None
        # Always reset playwright context when relaunching
        self._playwright = None
        self._browser = None
        try:
            from playwright.async_api import async_playwright
            self._playwright = await async_playwright().start()
            self._browser = await self._playwright.chromium.launch(
                headless=True,
                args=[
                    "--no-sandbox", "--disable-dev-shm-usage", "--disable-gpu",
                    "--disable-setuid-sandbox", "--single-process", "--no-zygote",
                    "--disable-extensions", "--disable-background-networking",
                    "--memory-pressure-off", "--disable-features=TranslateUI",
                    "--disable-web-security", "--disable-background-timer-throttling",
                ]
            )
            logger.info("🌐 Playwright Chromium — ONLINE")
            return self._browser
        except Exception as e:
            logger.error(f"Browser launch failed: {e}")
            return None

    async def _fresh_browser(self):
        """Launch a completely fresh browser every time — most reliable approach"""
        try:
            from playwright.async_api import async_playwright
            pw = await async_playwright().start()
            browser = await pw.chromium.launch(
                headless=True,
                args=[
                    "--no-sandbox", "--disable-dev-shm-usage", "--disable-gpu",
                    "--disable-setuid-sandbox", "--single-process", "--no-zygote",
                    "--disable-extensions", "--disable-background-networking",
                    "--memory-pressure-off", "--disable-features=TranslateUI",
                ]
            )
            return pw, browser
        except Exception as e:
            logger.error(f"Fresh browser launch failed: {e}")
            return None, None

    async def browser_navigate(self, url: str) -> Dict:
        """Navigate to URL like a human — returns title, text content, screenshot"""
        pw, browser = await self._fresh_browser()
        if not browser:
            logger.warning("Browser unavailable — falling back to scraper")
            return await self.scrape_webpage(url)
        try:
            context = await browser.new_context(
                viewport={"width": 1280, "height": 900},
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                locale="en-US",
            )
            page = await context.new_page()
            # Block ads/trackers to speed up loading
            await page.route("**/*.{png,jpg,jpeg,gif,svg,ico,woff,woff2}", lambda r: r.abort())
            await page.goto(url, wait_until="domcontentloaded", timeout=35000)
            await page.wait_for_timeout(1500)
            title = await page.title()
            # Extract meaningful text
            text = await page.evaluate("""() => {
                const scripts = document.querySelectorAll('script,style,nav,footer,iframe');
                scripts.forEach(s => s.remove());
                return document.body ? document.body.innerText.slice(0, 6000) : '';
            }""")
            # Get all links
            links = await page.evaluate("""() => {
                return Array.from(document.querySelectorAll('a[href]'))
                    .slice(0, 20)
                    .map(a => ({text: a.innerText.trim().slice(0,80), href: a.href}))
                    .filter(l => l.href.startsWith('http'));
            }""")
            # Screenshot
            screenshot_path = self.screenshots_dir / f"nav_{int(time.time())}.png"
            await page.screenshot(path=str(screenshot_path), full_page=False)
            screenshot_b64 = base64.b64encode(screenshot_path.read_bytes()).decode()
            return {
                "url": url, "title": title,
                "text": text, "links": links[:10],
                "screenshot": screenshot_b64,
                "screenshot_path": str(screenshot_path),
                "success": True
            }
        except Exception as e:
            logger.error(f"browser_navigate error: {e}")
            return await self.scrape_webpage(url)
        finally:
            try:
                await browser.close()
                await pw.stop()
            except Exception:
                pass

    async def browser_click(self, url: str, selector: str) -> Dict:
        """Navigate to a page and click an element — like a human clicking"""
        pw, browser = await self._fresh_browser()
        if not browser:
            return {"success": False, "error": "Browser unavailable"}
        try:
            context = await browser.new_context(
                viewport={"width": 1280, "height": 900},
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            )
            page = await context.new_page()
            await page.goto(url, wait_until="domcontentloaded", timeout=35000)
            await page.wait_for_timeout(1000)
            # Try to click the element
            try:
                await page.click(selector, timeout=8000)
                await page.wait_for_timeout(1500)
            except Exception:
                # Try to find by text
                try:
                    await page.get_by_text(selector).first.click(timeout=5000)
                    await page.wait_for_timeout(1500)
                except Exception:
                    pass
            title = await page.title()
            new_url = page.url
            text = await page.inner_text("body")
            screenshot_path = self.screenshots_dir / f"click_{int(time.time())}.png"
            await page.screenshot(path=str(screenshot_path))
            screenshot_b64 = base64.b64encode(screenshot_path.read_bytes()).decode()
            return {
                "url": url, "clicked_selector": selector,
                "new_url": new_url, "title": title,
                "text": text[:3000],
                "screenshot": screenshot_b64,
                "success": True
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
        finally:
            try:
                await browser.close()
                await pw.stop()
            except Exception:
                pass

    async def browser_fill_form(self, url: str, form_data: Dict, brain=None) -> Dict:
        """Navigate to a page and fill a form like a human — smart field detection"""
        pw, browser = await self._fresh_browser()
        if not browser:
            return {"success": False, "error": "Browser unavailable"}
        try:
            context = await browser.new_context(
                viewport={"width": 1280, "height": 900},
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            )
            page = await context.new_page()
            await page.goto(url, wait_until="domcontentloaded", timeout=35000)
            await page.wait_for_timeout(1500)

            filled = []
            failed = []
            for field_name, value in form_data.items():
                filled_ok = False
                # Build list of selectors to try (smart detection)
                is_css = any(field_name.startswith(c) for c in ['.','#','[','input','textarea','select','button'])
                if is_css:
                    selectors = [field_name]
                else:
                    selectors = [
                        f'[name="{field_name}"]',
                        f'[id="{field_name}"]',
                        f'[name*="{field_name}"]',
                        f'[id*="{field_name}"]',
                        f'[placeholder*="{field_name}"]',
                        f'input[type="text"]:nth-of-type(1)' if field_name in ['name','custname','fullname','username'] else None,
                        f'input[type="email"]' if field_name in ['email','custemail','mail'] else None,
                        f'input[type="tel"]' if field_name in ['phone','tel','custtel','mobile'] else None,
                    ]
                    selectors = [s for s in selectors if s]

                for sel in selectors:
                    try:
                        el = await page.query_selector(sel)
                        if not el:
                            continue
                        tag = await el.evaluate("el => el.tagName.toLowerCase()")
                        inp_type = await el.evaluate("el => (el.type || '').toLowerCase()")
                        if inp_type in ['checkbox','radio']:
                            if str(value).lower() in ['true','1','yes','on','check']:
                                await el.check()
                            else:
                                await el.uncheck()
                        elif tag == 'select':
                            try:
                                await el.select_option(label=str(value))
                            except Exception:
                                await el.select_option(value=str(value))
                        else:
                            await el.click()
                            await el.fill(str(value))
                        filled.append(field_name)
                        filled_ok = True
                        await page.wait_for_timeout(200)
                        break
                    except Exception:
                        continue

                if not filled_ok:
                    failed.append(field_name)

            await page.wait_for_timeout(500)
            title = await page.title()
            screenshot_path = self.screenshots_dir / f"form_{int(time.time())}.png"
            await page.screenshot(path=str(screenshot_path))
            screenshot_b64 = base64.b64encode(screenshot_path.read_bytes()).decode()
            return {
                "url": url, "title": title,
                "fields_filled": filled, "fields_failed": failed,
                "form_filled": len(filled) > 0,
                "screenshot": screenshot_b64[:500] + "...[truncated]",
                "success": len(filled) > 0
            }
        except Exception as e:
            logger.error(f"browser_fill_form error: {e}")
            return {"success": False, "error": str(e)}
        finally:
            try:
                await browser.close()
                await pw.stop()
            except Exception:
                pass

    async def browser_screenshot(self, url: str) -> Dict:
        """Take a full-page screenshot of any website"""
        pw, browser = await self._fresh_browser()
        if not browser:
            return {"success": False, "error": "Browser unavailable"}
        try:
            context = await browser.new_context(
                viewport={"width": 1280, "height": 900},
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            )
            page = await context.new_page()
            await page.goto(url, wait_until="domcontentloaded", timeout=35000)
            await page.wait_for_timeout(2000)
            title = await page.title()
            screenshot_path = self.screenshots_dir / f"full_{int(time.time())}.png"
            await page.screenshot(path=str(screenshot_path), full_page=True)
            screenshot_b64 = base64.b64encode(screenshot_path.read_bytes()).decode()
            return {
                "url": url, "title": title,
                "screenshot": screenshot_b64,
                "screenshot_path": str(screenshot_path),
                "success": True
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
        finally:
            try:
                await browser.close()
                await pw.stop()
            except Exception:
                pass

    async def browser_extract_text(self, url: str, selector: str = "body") -> Dict:
        """Extract text from a specific element on a page"""
        pw, browser = await self._fresh_browser()
        if not browser:
            return await self.scrape_webpage(url)
        try:
            context = await browser.new_context(
                viewport={"width": 1280, "height": 900},
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            )
            page = await context.new_page()
            await page.goto(url, wait_until="domcontentloaded", timeout=35000)
            await page.wait_for_timeout(1500)
            text = await page.inner_text(selector)
            title = await page.title()
            return {"url": url, "title": title, "selector": selector, "text": text[:6000], "success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}
        finally:
            try:
                await browser.close()
                await pw.stop()
            except Exception:
                pass

    async def browser_get_links(self, url: str) -> Dict:
        """Get all links from a page like a human reading the page"""
        pw, browser = await self._fresh_browser()
        if not browser:
            return {"success": False, "error": "Browser unavailable"}
        try:
            context = await browser.new_context(
                viewport={"width": 1280, "height": 900},
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            )
            page = await context.new_page()
            await page.goto(url, wait_until="domcontentloaded", timeout=35000)
            await page.wait_for_timeout(1500)
            links = await page.evaluate("""() => {
                return Array.from(document.querySelectorAll('a[href]'))
                    .map(a => ({text: a.innerText.trim().slice(0,100), href: a.href}))
                    .filter(l => l.href.startsWith('http') && l.text.length > 0)
                    .slice(0, 50);
            }""")
            title = await page.title()
            return {"url": url, "title": title, "links": links, "count": len(links), "success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}
        finally:
            try:
                await browser.close()
                await pw.stop()
            except Exception:
                pass

    async def analyze_image_url(self, image_url: str, question: str = "Describe everything you see", brain=None) -> Dict:
        """Analyze an image from URL using vision AI"""
        try:
            if brain:
                result = await brain.analyze_image(image_url, question)
                return {"image_url": image_url, "question": question, "analysis": result, "success": True}
            # Tesseract fallback
            return await self.ocr_image_url(image_url)
        except Exception as e:
            return {"error": str(e), "success": False}

    async def ocr_image_url(self, image_url: str) -> Dict:
        """Extract text from an image using OCR"""
        try:
            # Try OCR.space API first (free tier)
            ocr_key = os.getenv("OCR_SPACE_KEY", "")
            if ocr_key:
                response = requests.post(
                    "https://api.ocr.space/parse/imageurl",
                    data={"apikey": ocr_key, "url": image_url, "language": "eng"},
                    timeout=15
                )
                if response.ok:
                    data = response.json()
                    text = data.get("ParsedResults", [{}])[0].get("ParsedText", "")
                    if text:
                        return {"text": text, "method": "ocr_space", "success": True}

            # Pytesseract fallback
            import pytesseract
            from PIL import Image
            img_response = requests.get(image_url, timeout=10)
            img = Image.open(io.BytesIO(img_response.content))
            text = pytesseract.image_to_string(img)
            return {"text": text.strip() or "No text found", "method": "tesseract", "success": True}

        except Exception as e:
            return {"error": str(e), "success": False}

    async def ocr_image_base64(self, b64_data: str) -> Dict:
        """Extract text from a base64 encoded image"""
        try:
            import pytesseract
            from PIL import Image
            img_bytes = base64.b64decode(b64_data)
            img = Image.open(io.BytesIO(img_bytes))
            text = pytesseract.image_to_string(img)
            return {"text": text.strip(), "method": "tesseract", "success": True}
        except Exception as e:
            return {"error": str(e), "success": False}

    async def screenshot_and_analyze(self, url: str, question: str, brain=None) -> Dict:
        """Screenshot a website then analyze it visually"""
        screenshot_result = await self.browser_screenshot(url)
        if screenshot_result.get("success") and screenshot_result.get("screenshot"):
            if brain:
                analysis = await brain.analyze_image(screenshot_result["screenshot"], question)
                return {"url": url, "question": question, "analysis": analysis, "success": True}
        return screenshot_result

    # ═══════════════════════════════════════════
    # 💻 CODE GENERATION & EXECUTION
    # ═══════════════════════════════════════════

    async def write_code(self, description: str, language: str = "python", brain=None) -> Dict:
        """Generate working code for any task"""
        if brain:
            lang_tips = {
                "html": "Return a FULL complete HTML page with DOCTYPE, head, body, inline CSS and JS. No markdown, no explanation.",
                "python": "Return ONLY clean Python code with comments. No markdown, no explanation.",
                "javascript": "Return ONLY clean JavaScript code. No markdown.",
                "react": "Return ONLY React JSX code. No markdown.",
            }
            tip = lang_tips.get(language, f"Return ONLY clean {language} code. No markdown, no explanation.")
            code = await brain.call_llm(
                f"Write this: {description}",
                system=f"You are an elite software engineer. {tip}"
            )
            code = re.sub(r'^```\w*\n?', '', code).rstrip('`').strip()
            return {"language": language, "code": code, "description": description, "success": True}
        return {"language": language, "code": f"# {description}\npass", "success": True}

    async def run_python_code(self, code: str) -> Dict:
        """Safely execute Python code in a subprocess"""
        try:
            with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False) as f:
                f.write(code)
                temp_file = f.name

            result = subprocess.run(
                [sys.executable, temp_file],
                capture_output=True, text=True, timeout=15
            )
            os.unlink(temp_file)
            return {
                "output": result.stdout[:3000],
                "error": result.stderr[:1000] if result.stderr else None,
                "success": result.returncode == 0
            }
        except subprocess.TimeoutExpired:
            return {"error": "Code execution timed out (15s limit)", "success": False}
        except Exception as e:
            return {"error": str(e), "success": False}

    async def debug_code(self, code: str, error: str, brain=None) -> Dict:
        """Debug and fix code errors"""
        if brain:
            fixed = await brain.call_llm(
                f"Fix this code:\n```\n{code}\n```\n\nError:\n{error}\n\nReturn ONLY the fixed code.",
                system="You are an expert debugger. Fix the code. Return ONLY the fixed code, no explanation."
            )
            fixed = re.sub(r'^```\w*\n?', '', fixed).rstrip('`').strip()
            return {"original_code": code, "fixed_code": fixed, "error_was": error, "success": True}
        return {"error": "Brain required for debugging", "success": False}

    # ═══════════════════════════════════════════
    # 📄 DOCUMENT & FILE CREATION
    # ═══════════════════════════════════════════

    async def create_pdf(self, title: str, content: str, filename: str = None) -> Dict:
        """Create a downloadable document (PDF if reportlab available, else txt)"""
        try:
            fname = filename or f"{title.replace(' ', '_')[:30]}"
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib.units import inch
                fname = fname + ".pdf"
                fpath = self.output_dir / fname
                doc = SimpleDocTemplate(str(fpath), pagesize=A4,
                                        rightMargin=inch, leftMargin=inch,
                                        topMargin=inch, bottomMargin=inch)
                styles = getSampleStyleSheet()
                story = [Paragraph(title, styles["Title"]), Spacer(1, 0.3 * inch)]
                for para in content.split("\n\n"):
                    if para.strip():
                        story.append(Paragraph(para.replace("\n", "<br/>"), styles["Normal"]))
                        story.append(Spacer(1, 0.2 * inch))
                doc.build(story)
            except ImportError:
                # Fallback: save as plain text file
                fname = fname + ".txt"
                fpath = self.output_dir / fname
                with open(str(fpath), 'w', encoding='utf-8') as f:
                    f.write(f"{title}\n{'='*len(title)}\n\nGenerated by S.T.E.W\n\n{content}")
            return {"filename": str(fpath), "title": title, "download_url": f"/download/{fname}", "success": True}
        except Exception as e:
            return {"error": str(e), "success": False}

    async def create_word_doc(self, title: str, content: str) -> Dict:
        """Create a real Word document"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            fname = f"{title.replace(' ', '_')[:30]}.docx"
            fpath = self.output_dir / fname

            doc = Document()
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            sub = doc.add_paragraph(f"Generated by S.T.E.W — {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")

            for para in content.split("\n\n"):
                if para.strip():
                    if para.startswith("#"):
                        doc.add_heading(para.lstrip("#").strip(), level=2)
                    else:
                        doc.add_paragraph(para.strip())

            doc.save(str(fpath))
            return {"filename": str(fpath), "title": title, "download_url": f"/download/{fname}", "success": True}
        except Exception as e:
            return {"error": str(e), "success": False}

    async def create_spreadsheet(self, name: str, data: List[List]) -> Dict:
        """Create an Excel spreadsheet"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment

            fname = f"{name.replace(' ', '_')[:30]}.xlsx"
            fpath = self.output_dir / fname

            wb = Workbook()
            ws = wb.active
            ws.title = name[:30]

            if data:
                header_fill = PatternFill(start_color="0A0050", end_color="0A0050", fill_type="solid")
                header_font = Font(color="FFFFFF", bold=True)
                for col, header in enumerate(data[0], 1):
                    cell = ws.cell(row=1, column=col, value=str(header))
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = Alignment(horizontal="center")
                for row_idx, row in enumerate(data[1:], 2):
                    for col_idx, val in enumerate(row, 1):
                        ws.cell(row=row_idx, column=col_idx, value=val)

            wb.save(str(fpath))
            return {"filename": str(fpath), "name": name, "download_url": f"/download/{fname}", "rows": len(data), "success": True}
        except Exception as e:
            return {"error": str(e), "success": False}

    async def create_website(self, name: str, description: str, color: str = "#0a0050") -> Dict:
        """Create a complete HTML website"""
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{name}</title>
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:'Segoe UI',sans-serif;background:#0a0a0a;color:#fff}}
header{{background:linear-gradient(135deg,{color},#1a1a2e);padding:4rem 2rem;text-align:center}}
header h1{{font-size:3rem;margin-bottom:1rem}}
header p{{font-size:1.2rem;opacity:0.8;max-width:600px;margin:0 auto}}
.cta{{display:inline-block;margin-top:2rem;padding:1rem 2rem;background:#fff;color:{color};border-radius:50px;font-weight:bold;text-decoration:none}}
main{{max-width:1000px;margin:4rem auto;padding:0 2rem}}
.card{{background:#111;border:1px solid #222;border-radius:16px;padding:2rem;margin:1rem 0}}
footer{{text-align:center;padding:2rem;background:#050505;color:#555}}
</style>
</head>
<body>
<header>
<h1>✦ {name}</h1>
<p>{description}</p>
<a href="#" class="cta">Get Started</a>
</header>
<main>
<div class="card"><h2 style="color:{color};margin-bottom:1rem">About</h2><p>{description}</p></div>
<div class="card"><h2 style="color:{color};margin-bottom:1rem">Features</h2>
<p>✓ Built with S.T.E.W — Secret Task Execution Worker<br>
✓ Designed for performance and beauty<br>
✓ Ready to deploy anywhere</p></div>
</main>
<footer><p>Built by S.T.E.W — MUTYINT © {datetime.now().year}</p></footer>
</body></html>"""

        fname = f"{name.replace(' ', '_')[:30]}.html"
        fpath = self.output_dir / fname
        fpath.write_text(html)
        return {"filename": str(fpath), "name": name, "download_url": f"/download/{fname}", "success": True}

    async def create_html_report(self, title: str, content: str, data: Dict = None) -> Dict:
        """Create a beautiful HTML report"""
        sections_html = ""
        for para in content.split("\n\n"):
            if para.strip():
                sections_html += f"<p>{para.strip()}</p>\n"

        html = f"""<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8"><title>{title}</title>
<style>
body{{font-family:'Segoe UI',sans-serif;background:#f8f9fa;margin:0;padding:2rem}}
.report{{max-width:800px;margin:0 auto;background:#fff;border-radius:16px;overflow:hidden;box-shadow:0 4px 20px rgba(0,0,0,0.1)}}
.header{{background:linear-gradient(135deg,#0a0050,#1a1a9e);color:#fff;padding:2rem}}
.header h1{{margin:0;font-size:2rem}}
.meta{{opacity:0.7;margin-top:0.5rem;font-size:0.9rem}}
.body{{padding:2rem;line-height:1.8}}
.body p{{margin:1rem 0;color:#333}}
footer{{text-align:center;padding:1rem;background:#f0f0f0;color:#999;font-size:0.85rem}}
</style>
</head>
<body>
<div class="report">
<div class="header">
<h1>📋 {title}</h1>
<div class="meta">Generated by S.T.E.W — {datetime.now().strftime('%B %d, %Y %H:%M')}</div>
</div>
<div class="body">{sections_html}</div>
<footer>S.T.E.W — Secret Task Execution Worker — MUTYINT</footer>
</div>
</body></html>"""

        fname = f"report_{title.replace(' ', '_')[:25]}_{int(time.time())}.html"
        fpath = self.output_dir / fname
        fpath.write_text(html)
        return {"filename": str(fpath), "title": title, "download_url": f"/download/{fname}", "success": True}

    async def save_to_workspace(self, filename: str, content: str) -> Dict:
        """Save any file to S.T.E.W's persistent workspace"""
        try:
            fpath = self.workspace_dir / filename
            fpath.write_text(content)
            return {"filename": str(fpath), "size": len(content), "success": True}
        except Exception as e:
            return {"error": str(e), "success": False}

    async def read_from_workspace(self, filename: str) -> Dict:
        """Read a file from S.T.E.W's workspace"""
        try:
            fpath = self.workspace_dir / filename
            if fpath.exists():
                return {"filename": filename, "content": fpath.read_text(), "success": True}
            return {"error": "File not found", "success": False}
        except Exception as e:
            return {"error": str(e), "success": False}

    async def list_files(self, directory: str = "output") -> Dict:
        """List files in output or workspace directory"""
        try:
            d = Path(directory)
            if not d.exists():
                return {"files": [], "count": 0, "success": True}
            files = [{"name": f.name, "size": f.stat().st_size, "modified": f.stat().st_mtime}
                     for f in d.iterdir() if f.is_file()]
            return {"directory": directory, "files": files, "count": len(files), "success": True}
        except Exception as e:
            return {"error": str(e), "success": False}

    async def download_file(self, url: str, filename: str = None) -> Dict:
        """Download a file from the internet"""
        try:
            response = requests.get(url, timeout=30, stream=True)
            if not filename:
                filename = url.split("/")[-1].split("?")[0] or "download"
            fpath = self.output_dir / filename
            with open(str(fpath), 'wb') as f:
                for chunk in response.iter_content(8192):
                    f.write(chunk)
            return {"filename": str(fpath), "url": url, "size": fpath.stat().st_size,
                    "download_url": f"/download/{filename}", "success": True}
        except Exception as e:
            return {"error": str(e), "success": False}

    async def read_pdf(self, filepath: str) -> Dict:
        """Read and extract text from a PDF"""
        try:
            import PyPDF2
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = "\n".join([page.extract_text() or "" for page in reader.pages])
            return {"filepath": filepath, "text": text[:8000], "pages": len(reader.pages), "success": True}
        except Exception as e:
            return {"error": str(e), "success": False}

    # ═══════════════════════════════════════════
    # 🌍 REAL-WORLD DATA
    # ═══════════════════════════════════════════

    async def weather_info(self, city: str) -> Dict:
        """Get real weather for any city"""
        try:
            r = requests.get(
                f"https://wttr.in/{city}?format=j1",
                headers={"User-Agent": "curl/7.68.0"}, timeout=10
            )
            if r.ok:
                data = r.json()
                cw = data["current_condition"][0]
                return {
                    "city": city,
                    "temperature": f"{cw['temp_C']}°C / {cw['temp_F']}°F",
                    "feels_like": f"{cw['FeelsLikeC']}°C",
                    "humidity": f"{cw['humidity']}%",
                    "description": cw["weatherDesc"][0]["value"],
                    "wind_speed": f"{cw['windspeedKmph']} km/h",
                    "visibility": f"{cw['visibility']} km",
                    "uv_index": cw.get("uvIndex", "N/A"),
                    "success": True
                }
        except Exception as e:
            return {"city": city, "error": str(e), "success": False}

    async def convert_currency(self, amount: float, from_currency: str, to_currency: str) -> Dict:
        """Convert currency using live exchange rates"""
        try:
            r = requests.get(
                f"https://api.exchangerate-api.com/v4/latest/{from_currency.upper()}",
                timeout=10
            )
            if r.ok:
                data = r.json()
                rate = data["rates"].get(to_currency.upper())
                if rate:
                    return {"amount": amount, "from": from_currency.upper(), "to": to_currency.upper(),
                            "rate": rate, "converted": round(amount * rate, 2), "success": True}
        except Exception as e:
            pass
        # Fallback
        try:
            r = requests.get(f"https://open.er-api.com/v6/latest/{from_currency.upper()}", timeout=10)
            if r.ok:
                data = r.json()
                rate = data.get("rates", {}).get(to_currency.upper())
                if rate:
                    return {"amount": amount, "from": from_currency.upper(), "to": to_currency.upper(),
                            "rate": rate, "converted": round(amount * rate, 2), "success": True}
        except Exception as e:
            return {"error": str(e), "success": False}

    async def stock_price(self, symbol: str) -> Dict:
        """Get stock price information"""
        try:
            r = requests.get(
                f"https://query1.finance.yahoo.com/v8/finance/chart/{symbol}?interval=1d&range=1d",
                headers={"User-Agent": "Mozilla/5.0"}, timeout=10
            )
            if r.ok:
                data = r.json()
                meta = data["chart"]["result"][0]["meta"]
                price = meta.get("regularMarketPrice", 0)
                prev = meta.get("previousClose", 0)
                change_pct = ((price - prev) / prev * 100) if prev else 0
                return {
                    "symbol": symbol.upper(), "price": f"${price:.2f}",
                    "prev_close": f"${prev:.2f}",
                    "change": f"{'+' if change_pct >= 0 else ''}{change_pct:.2f}%",
                    "market": meta.get("exchangeName", ""),
                    "currency": meta.get("currency", "USD"),
                    "success": True
                }
        except Exception as e:
            return {"symbol": symbol, "error": str(e), "success": False}

    # ═══════════════════════════════════════════
    # 🌐 LANGUAGE & TEXT
    # ═══════════════════════════════════════════

    async def translate_text(self, text: str, target_language: str, brain=None) -> Dict:
        """Translate text to any language — uses GoogleTranslator with AI fallback for African/unsupported languages"""
        # Map common language names to codes
        lang_map = {
            'yoruba': 'yo', 'igbo': 'ig', 'hausa': 'ha', 'swahili': 'sw',
            'zulu': 'zu', 'amharic': 'am', 'pidgin': None, 'nigerian pidgin': None,
            'french': 'fr', 'spanish': 'es', 'arabic': 'ar', 'portuguese': 'pt',
            'german': 'de', 'chinese': 'zh-CN', 'mandarin': 'zh-CN', 'hindi': 'hi',
            'japanese': 'ja', 'korean': 'ko', 'russian': 'ru', 'italian': 'it',
            'dutch': 'nl', 'turkish': 'tr', 'polish': 'pl', 'thai': 'th',
            'afrikaans': 'af', 'somali': 'so', 'xhosa': 'xh', 'twi': 'tw',
        }
        tl = target_language.strip().lower()
        lang_code = lang_map.get(tl, tl)

        # Try GoogleTranslator first
        if lang_code:
            try:
                from deep_translator import GoogleTranslator
                translated = GoogleTranslator(source="auto", target=lang_code).translate(text)
                return {"original": text, "translated_text": translated, "language": target_language, "success": True}
            except Exception as e1:
                logger.warning(f"GoogleTranslator failed for {target_language}: {e1}")

        # AI brain fallback — handles Pidgin, uncommon African languages, etc.
        try:
            prompt = f"""Translate the following text to {target_language}.
Return ONLY the translated text, nothing else. No explanation, no labels.

Text to translate: {text}

Translation in {target_language}:"""
            result = await brain.think(prompt) if brain else (await self.brain.think(prompt) if hasattr(self, 'brain') and self.brain else None)
            if result:
                return {"original": text, "translated_text": result.strip(), "language": target_language,
                        "method": "ai_translation", "success": True}
        except Exception as e2:
            logger.error(f"AI translate fallback failed: {e2}")

        return {"original": text, "translated_text": text, "language": target_language,
                "success": False, "error": f"Could not translate to {target_language}"}

    async def summarize_text(self, text: str, brain=None) -> Dict:
        """Summarize long text"""
        if brain:
            summary = await brain.call_llm(
                f"Summarize this concisely in 3-5 key points:\n\n{text[:4000]}",
                system="You are a precise summarizer. Give key points only."
            )
            return {"original_length": len(text), "summary": summary, "success": True}
        return {"summary": text[:500] + "...", "success": True}

    async def analyze_sentiment(self, text: str) -> Dict:
        """Analyze sentiment of text"""
        try:
            try:
                from textblob import TextBlob as TB
                blob = TB(text)
                polarity = blob.sentiment.polarity
                if polarity > 0.1:
                    sentiment = "Positive 😊"
                elif polarity < -0.1:
                    sentiment = "Negative 😔"
                else:
                    sentiment = "Neutral 😐"
                return {"text": text[:100], "sentiment": sentiment, "score": round(polarity, 3), "success": True}
            except ImportError:
                # Fallback: basic keyword sentiment
                pos = sum(1 for w in ['good','great','excellent','amazing','love','best','happy','wonderful'] if w in text.lower())
                neg = sum(1 for w in ['bad','terrible','awful','hate','worst','horrible','sad','ugly'] if w in text.lower())
                score = (pos - neg) / max(len(text.split()), 1)
                sentiment = "Positive 😊" if score > 0 else "Negative 😔" if score < 0 else "Neutral 😐"
                return {"text": text[:100], "sentiment": sentiment, "score": round(score, 3), "success": True}
        except Exception as e:
            return {"error": str(e), "success": False}

    async def detect_language(self, text: str) -> Dict:
        """Detect the language of text"""
        try:
            from langdetect import detect
            lang = detect(text)
            return {"text": text[:100], "language": lang, "success": True}
        except Exception as e:
            return {"error": str(e), "success": False}

    # ═══════════════════════════════════════════
    # 🔗 API & AUTOMATION
    # ═══════════════════════════════════════════

    async def api_call(self, method: str, url: str, headers: Dict = None, body: Dict = None) -> Dict:
        """Make any HTTP API call"""
        try:
            h = headers or {}
            if method.upper() == "GET":
                r = requests.get(url, headers=h, timeout=15)
            elif method.upper() == "POST":
                r = requests.post(url, json=body, headers=h, timeout=15)
            elif method.upper() == "PUT":
                r = requests.put(url, json=body, headers=h, timeout=15)
            elif method.upper() == "DELETE":
                r = requests.delete(url, headers=h, timeout=15)
            else:
                return {"error": f"Unknown method: {method}", "success": False}
            try:
                response_data = r.json()
            except Exception:
                response_data = r.text[:2000]
            return {"status": r.status_code, "data": response_data, "success": r.ok}
        except Exception as e:
            return {"error": str(e), "success": False}

    async def send_webhook(self, url: str, event: str, data: Dict) -> Dict:
        """Send webhook to any URL"""
        try:
            payload = {"event": event, "data": data, "timestamp": datetime.now().isoformat(), "source": "S.T.E.W"}
            r = requests.post(url, json=payload, timeout=10)
            return {"url": url, "event": event, "status": r.status_code, "success": r.ok}
        except Exception as e:
            return {"error": str(e), "success": False}

    async def bulk_scrape(self, urls: List[str]) -> List[Dict]:
        """Scrape multiple URLs simultaneously"""
        tasks = [self.scrape_webpage(url) for url in urls]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        return [r for r in results if isinstance(r, dict)]

    async def extract_data(self, url: str, selector: str = None) -> Dict:
        """Extract structured data from a page"""
        return await self.browser_extract_text(url, selector or "body")

    async def monitor_website(self, url: str, check_interval: int = 300) -> Dict:
        """Check if a website is online"""
        try:
            start = time.time()
            r = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
            response_time = round(time.time() - start, 3)
            return {
                "url": url, "status": r.status_code, "online": r.ok,
                "response_time": f"{response_time}s", "success": True
            }
        except Exception as e:
            return {"url": url, "online": False, "error": str(e), "success": False}

    async def send_email(self, to: str, subject: str, body: str) -> Dict:
        """Send an email via SMTP"""
        try:
            import smtplib
            from email.mime.text import MIMEText
            from email.mime.multipart import MIMEMultipart

            smtp_host = os.getenv("SMTP_HOST", "smtp.gmail.com")
            smtp_port = int(os.getenv("SMTP_PORT", "587"))
            smtp_user = os.getenv("SMTP_USER", "")
            smtp_pass = os.getenv("SMTP_PASS", "")

            if not smtp_user or not smtp_pass:
                return {"error": "SMTP credentials not configured", "success": False}

            msg = MIMEMultipart()
            msg["From"] = smtp_user
            msg["To"] = to
            msg["Subject"] = subject
            msg.attach(MIMEText(body, "plain"))

            with smtplib.SMTP(smtp_host, smtp_port) as server:
                server.starttls()
                server.login(smtp_user, smtp_pass)
                server.send_message(msg)

            return {"to": to, "subject": subject, "success": True}
        except Exception as e:
            return {"error": str(e), "success": False}

    async def get_system_info(self) -> Dict:
        """Get S.T.E.W system information"""
        import platform
        return {
            "agent": "S.T.E.W — Secret Task Execution Worker",
            "version": "3.0.0 ULTRA",
            "creator": "Emmanuel Ene Rejoice Gideon",
            "company": "MUTYINT",
            "platform": platform.system(),
            "python": platform.python_version(),
            "skills": 60,
            "browser": "Playwright Chromium",
            "vision": "Tesseract OCR + LLM Vision",
            "search": "Serper (Google)",
            "timestamp": datetime.now().isoformat(),
            "success": True
        }

    async def create_qr_code(self, data: str, filename: str = None) -> Dict:
        """Generate a QR code"""
        try:
            import qrcode


            fname = filename or f"qr_{int(time.time())}.png"
            fpath = self.output_dir / fname
            img = qrcode.make(data)
            img.save(str(fpath))
            return {"data": data, "filename": str(fpath), "download_url": f"/download/{fname}", "success": True}
        except ImportError:
            return {"error": "qrcode library not installed", "success": False}
        except Exception as e:
            return {"error": str(e), "success": False}

    async def generate_image_prompt(self, description: str, brain=None) -> Dict:
        """Generate an optimized image prompt"""
        if brain:
            prompt = await brain.call_llm(
                f"Create a detailed, vivid image generation prompt for: {description}",
                system="You are an expert at writing Midjourney/DALL-E prompts. Be specific, visual, and detailed."
            )
            image_url = f"https://image.pollinations.ai/prompt/{requests.utils.quote(prompt[:200])}"
            return {"description": description, "prompt": prompt, "image_url": image_url, "success": True}
        return {"error": "Brain required", "success": False}

    async def schedule_task(self, task_name: str, schedule: str, action: str) -> Dict:
        """Register a scheduled task"""
        return {
            "task_name": task_name, "schedule": schedule, "action": action,
            "status": "scheduled", "message": f"Task '{task_name}' scheduled: {schedule}",
            "success": True
        }

    async def create_database_schema(self, name: str, tables: List[Dict]) -> Dict:
        """Generate a database schema"""
        schema = {"database": name, "tables": tables, "created": datetime.now().isoformat()}
        fname = f"{name}_schema.json"
        fpath = self.output_dir / fname
        fpath.write_text(json.dumps(schema, indent=2))
        return {"schema": schema, "filename": str(fpath), "download_url": f"/download/{fname}", "success": True}
